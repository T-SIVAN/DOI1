from __future__ import annotations

import json
import base64
import hashlib
import os
import re
import time
import urllib.error
import urllib.parse
import urllib.request
from io import BytesIO
from dataclasses import replace
from datetime import datetime, timezone
from typing import Any, Dict, List, Optional

import pandas as pd
import streamlit as st

from attachment_utils import SUPPORTED_ATTACHMENT_TYPES, build_attachment_context
from literature_search import (
    FederatedSearchResult,
    PaperRecord,
    export_bibtex,
    export_excel,
    export_ris,
    search_literature,
)
from ppt_report import build_pptx_bytes
from auth_service import AuthIdentity, require_authenticated_user
from research_store import ConfigurationError, ResearchStore, create_research_store
from research_workspace import render_research_conversation_sidebar, render_research_workspace


OPENALEX_WORKS = "https://api.openalex.org/works"
DEFAULT_OPENALEX_EMAIL = ""
DEFAULT_OPENALEX_API_KEY = ""
DEFAULT_LLM_BASE_URL = "https://api.openai.com/v1"
DEFAULT_LLM_MODEL = "gpt-4o-mini"
DEFAULT_GEMINI_BASE_URL = "https://generativelanguage.googleapis.com/v1beta"
DEFAULT_GEMINI_MODEL = "gemini-2.5-flash-lite"
GEMINI_FALLBACK_MODELS = [
    "gemini-2.5-flash-lite",
    "gemini-2.0-flash-lite",
    "gemini-2.0-flash",
    "gemini-1.5-flash",
    "gemini-2.5-flash",
]
PDF_TEXT_MODE = "文本模式（默认，速度快）"
PDF_IMAGE_MODE = "图片模式（适合扫描版/图表多）"
PDF_HYBRID_MODE = "混合模式（先文字，失败或内容太少时转图片）"
PDF_PARSE_FAST = "快速模式"
PDF_PARSE_ENHANCED = "增强模式"
PDF_PARSE_STRUCTURED = "结构化模式"
CITATION_METADATA_ONLY = "仅元数据"
CITATION_CONTEXT_MODE = "尝试开放全文上下文"
PDF_HYBRID_MIN_TEXT_CHARS = 800
PDF_IMAGE_RENDER_SCALE = 1.6
PDF_IMAGE_JPEG_QUALITY = 72
PDF_IMAGE_MAX_BYTES = 1_200_000
PPT_ANALYSIS_MAX_CHARS = 18000
JOURNAL_IMPACT_FACTORS = {
    "nature biotechnology": {
        "if": "44.5",
        "note": "2025 JCR / Nature Journal Metrics",
    },
    "nat biotechnol": {
        "if": "44.5",
        "note": "2025 JCR / Nature Journal Metrics",
    },
}


@st.cache_resource(show_spinner=False)
def cached_research_store() -> ResearchStore:
    """One shared connection pool/object client per Streamlit process."""
    return create_research_store()

OUTPUT_COLUMNS = [
    "文献标题",
    "年份",
    "DOI",
    "摘要",
    "OpenAlex ID",
    "引用数",
    "来源期刊",
    "作者",
    "证据等级",
    "引用上下文",
    "上下文来源",
    "PMID",
    "MeSH",
    "开放全文URL",
    "解析状态",
]

BREAKTHROUGH_COLUMNS = ["年份", "里程碑/技术突破", "代表文献", "DOI", "关键意义"]

NATURE_TOOL_CONFIGS = [
    {
        "category": "写作润色",
        "name": "Nature风格润色/翻译",
        "help": "将中文或英文论文段落改写为更接近 Nature / 高影响力期刊的学术表达。",
        "system": "你是熟悉 Nature 系列论文语言风格的资深学术英文编辑。",
        "instruction": """
请对输入文本进行学术润色、重构或中译英。输出：
1. 【润色后文本】：给出可直接粘贴到论文中的版本。
2. 【关键修改说明】：说明逻辑、术语、句式和语气的主要改动。
3. 【风险提示】：指出证据不足、表述过强或需要补充引用的位置。
""",
    },
    {
        "category": "写作润色",
        "name": "论文写作/章节起草",
        "help": "基于研究结果、笔记或要点起草摘要、引言、结果、讨论等章节。",
        "system": "你是擅长 Nature 风格论文论证结构设计的科研写作顾问。",
        "instruction": """
请根据输入材料起草或重建论文段落。输出：
1. 【建议标题/中心论点】。
2. 【章节正文草稿】：结构清晰，避免空泛。
3. 【论证链条】：用要点说明背景、缺口、方法、结果和意义如何连接。
4. 【还需补充的信息】：列出缺失数据或需要引用支撑的地方。
""",
    },
    {
        "category": "投稿审稿",
        "name": "预投稿审稿模拟",
        "help": "从审稿人视角评估创新性、技术可靠性、逻辑漏洞和补实验建议。",
        "system": "你是严格但建设性的高影响力期刊审稿人。",
        "instruction": """
请模拟预投稿审稿。输出：
1. 【总体判断】：接收潜力、主要短板和适合期刊层级。
2. 【Reviewer 1】：聚焦创新性和重要性。
3. 【Reviewer 2】：聚焦方法、统计和实验设计。
4. 【Reviewer 3】：聚焦表述、结构和可重复性。
5. 【优先修改清单】：按重要性排序。
""",
    },
    {
        "category": "文献数据",
        "name": "引用支撑与检索策略",
        "help": "把论文陈述拆成需要引用支撑的句子，并给出检索式、候选文献方向和 DOI 核验提示。",
        "system": "你是熟悉 Nature/CNS 文献引用规范和生物医药检索策略的文献顾问。",
        "instruction": """
请为输入文本规划引用支撑。当前网页不会实时联网检索，因此不要编造 DOI。输出：
1. 【需要引用支撑的陈述】：逐条列出。
2. 【推荐检索式】：给出 PubMed / CrossRef / Google Scholar 可用关键词。
3. 【候选文献类型】：说明应找原创研究、方法学论文还是综述。
4. 【已知候选 DOI】：只有在非常确定时才给出；不确定写“需核验”。
""",
    },
    {
        "category": "文献数据",
        "name": "数据可用性/FAIR声明",
        "help": "生成 Data Availability statement、数据仓储建议和 FAIR 检查清单。",
        "system": "你是熟悉高影响力期刊数据共享规范、FAIR 原则和科研数据仓储的编辑。",
        "instruction": """
请根据输入研究内容生成数据共享方案。输出：
1. 【Data Availability Statement】英文正式版本。
2. 【中文说明】：解释哪些数据需要公开、哪些可受限。
3. 【推荐仓储】：按数据类型推荐 GEO、SRA、PRIDE、Zenodo、Figshare 等。
4. 【FAIR 检查清单】：Findable、Accessible、Interoperable、Reusable。
""",
    },
    {
        "category": "投稿审稿",
        "name": "审稿意见回复",
        "help": "把 reviewer comments 转成逐点回复信，包含礼貌回应、修改动作和证据边界。",
        "system": "你是擅长撰写高质量 response to reviewers 的科研通讯作者。",
        "instruction": """
请为输入审稿意见起草逐点回复。输出：
1. 【总回复开头】。
2. 【逐点回复表】：Reviewer comment / Response / Manuscript change。
3. 【需要补实验或补分析的事项】。
4. 【语气风险】：指出过度承诺或回应不足的位置。
""",
    },
    {
        "category": "文献数据",
        "name": "论文精读/中英对照提纲",
        "help": "把论文片段整理成中文精读笔记、英文关键句和图文对应阅读提纲。",
        "system": "你是擅长论文精读、图文逻辑拆解和中英双语教学的科研导师。",
        "instruction": """
请对输入论文片段做精读。输出：
1. 【一句话结论】。
2. 【中文精读笔记】：背景、问题、方法、结果、意义。
3. 【英文关键句】：提炼 5-8 句关键英文表达。
4. 【图表阅读线索】：如果输入提到图表，请说明每张图应回答什么问题。
""",
    },
    {
        "category": "展示转化",
        "name": "论文汇报PPT大纲",
        "help": "生成组会/文献汇报用中文 PPT 结构、每页要点和讲稿提示。",
        "system": "你是擅长科研论文汇报和中文学术演示设计的导师。",
        "instruction": """
请把输入内容设计成文献汇报 PPT 大纲。输出：
1. 【汇报主线】。
2. 【逐页幻灯片大纲】：页码、标题、核心图/表、页面要点。
3. 【讲稿提示】：每页 1-3 句。
4. 【听众可能提问】：列出可能问题和回答思路。
""",
    },
    {
        "category": "展示转化",
        "name": "论文转专利初筛",
        "help": "从论文或技术方案中提取可专利化创新点、权利要求雏形和证据映射。",
        "system": "你是熟悉中国发明专利撰写和生物医药技术转化的专利工程师。",
        "instruction": """
请进行专利转化初筛。输出：
1. 【可专利化技术点】。
2. 【可能的独立权利要求雏形】。
3. 【从属权利要求方向】。
4. 【说明书证据映射】：每个技术特征对应输入材料中的依据。
5. 【新颖性/创造性风险】。
""",
    },
    {
        "category": "展示转化",
        "name": "科研绘图规划",
        "help": "为论文结果设计多面板 figure 逻辑、统计图类型和 Python/R 绘图建议。",
        "system": "你是擅长高影响力期刊科研图设计和数据可视化的 figure editor。",
        "instruction": """
请为输入结果设计投稿级科研图。输出：
1. 【图的核心结论】。
2. 【多面板设计】：Panel A/B/C... 每个面板展示什么。
3. 【推荐图型和统计标注】。
4. 【配色和版式建议】：保持克制、清晰、可发表。
5. 【需要的数据表结构】。
""",
    },
    {
        "category": "文献数据",
        "name": "多源文献检索方案",
        "help": "为一个研究问题生成 PubMed/CrossRef/arXiv/Scopus 检索词、筛选标准和导出字段。",
        "system": "你是熟悉多数据库文献检索、MeSH 词和引用管理的医学文献检索专家。",
        "instruction": """
请为输入研究问题设计文献检索方案。输出：
1. 【研究问题拆解】：PICO/关键词/同义词。
2. 【PubMed 检索式】：包含 MeSH 和自由词。
3. 【其他数据库检索式】：CrossRef、Google Scholar、Scopus 或 arXiv。
4. 【纳入排除标准】。
5. 【建议导出字段】：题名、摘要、DOI、PMID、年份、期刊等。
""",
    },
    {
        "category": "写作润色",
        "name": "基金申请书策划",
        "help": "把研究基础和初步结果组织为立项依据、科学问题、研究目标与技术路线。",
        "system": "你是熟悉生物医学基金申请评审标准的科研项目顾问。",
        "instruction": """
请根据附件和输入信息规划基金申请书。输出：
1. 【核心科学问题与假设】。
2. 【立项依据框架】：现状、缺口、前期基础和必要性。
3. 【研究目标与工作包】：目标、方法、里程碑和交付物。
4. 【技术路线与风险预案】。
5. 【仍需作者补充的证据】：不得代填未知数据。
""",
    },
    {
        "category": "文献数据",
        "name": "参考文献一致性核验",
        "help": "核对正文陈述与参考文献是否匹配，并标出 DOI、年份、题名等待查项。",
        "system": "你是严谨的科研参考文献审校员，只依据用户提供的正文与参考文献判断。",
        "instruction": """
请逐条核验正文陈述与所附参考文献。输出：
1. 【核验表】：陈述、对应文献、支持程度、问题和建议动作。
2. 【元数据异常】：题名、作者、年份、期刊、DOI 的矛盾或缺失。
3. 【过度外推】：文献结论不足以支持正文的地方。
4. 【待联网复核项】：无法从附件确认时明确标记，禁止编造。
""",
    },
    {
        "category": "研究设计",
        "name": "统计方案与功效分析",
        "help": "根据研究设计和数据结构选择统计方法，规划效应量、功效与敏感性分析。",
        "system": "你是生物统计师，重视研究设计、假设前提、效应量和可重复分析。",
        "instruction": """
请生成可执行的统计分析方案。输出：
1. 【变量与研究设计识别】。
2. 【主要/次要终点与统计模型】。
3. 【模型前提、缺失值、多重比较和敏感性分析】。
4. 【效应量与样本量/功效分析所需参数】；缺少参数时给公式和情景，不虚构数值。
5. 【推荐结果表、图和报告句式】。
""",
    },
    {
        "category": "研究设计",
        "name": "实验日志结构化",
        "help": "把实验记录、仪器导出或零散笔记整理为可追溯的实验日志和复现实验清单。",
        "system": "你是重视可追溯性、版本和偏差记录的实验室数据管理员。",
        "instruction": """
请把输入材料整理为结构化实验日志。输出：
1. 【实验目的、日期、人员与样本批次】。
2. 【材料、仪器、软件和版本】。
3. 【逐步方法与关键参数】。
4. 【原始观察、偏差、失败与处理决定】。
5. 【结果文件索引、质控检查和下一步】；未知字段保留“待补充”。
""",
    },
]


class ApiError(Exception):
    """Readable wrapper for recoverable HTTP/API failures."""


def normalize_identifier(value: str) -> str:
    return (value or "").strip().rstrip(".,;")


def normalize_doi(value: str) -> str:
    cleaned = normalize_identifier(value)
    cleaned = re.sub(r"^https?://(dx\.)?doi\.org/", "", cleaned, flags=re.I)
    cleaned = cleaned.strip().rstrip(").,;]")
    return cleaned


def extract_doi_from_text(text: Any) -> str:
    match = re.search(r"(?i)\b10\.\d{4,9}/[-._;()/:A-Z0-9]+", str(text or ""))
    return normalize_doi(match.group(0)) if match else ""


def normalize_journal_name(value: Any) -> str:
    text = strip_markup(value).lower()
    text = re.sub(r"[^a-z0-9& ]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def lookup_impact_factor(journal_name: Any) -> Dict[str, str]:
    normalized = normalize_journal_name(journal_name)
    if not normalized:
        return {"if": "待核验", "note": "未识别期刊名"}

    for key, value in JOURNAL_IMPACT_FACTORS.items():
        key_norm = normalize_journal_name(key)
        if normalized == key_norm or key_norm in normalized or normalized in key_norm:
            return value
    return {"if": "待核验", "note": "未在本地 IF 映射表中找到，请按 JCR/期刊官网核验"}


def is_openalex_id(value: str) -> bool:
    return bool(re.match(r"^(https://openalex\.org/)?W\d+$", value.strip(), flags=re.I))


def clean_openalex_id(value: str) -> str:
    return value.rstrip("/").split("/")[-1]


def strip_markup(text: Any) -> str:
    if text is None:
        return ""
    return re.sub(r"\s+", " ", re.sub(r"<[^>]+>", " ", str(text))).strip()


def compact_text(text: Any, max_chars: int = 7000) -> str:
    cleaned = re.sub(r"\s+", " ", str(text or "")).strip()
    if len(cleaned) <= max_chars:
        return cleaned
    return cleaned[: max_chars - 20].rstrip() + " ... [truncated]"


def config_value(name: str, default: str = "") -> str:
    """Read deploy-time configuration without requiring a local secrets file."""
    environment_value = os.getenv(name)
    if environment_value is not None:
        return environment_value
    try:
        value = st.secrets.get(name, default)
    except Exception:
        value = default
    return str(value) if value is not None else default


def load_config_defaults(llm_provider: str) -> Dict[str, Any]:
    contact_email = (
        config_value("LITERATURE_CONTACT_EMAIL")
        or config_value("NCBI_EMAIL")
        or config_value("OPENALEX_MAILTO", DEFAULT_OPENALEX_EMAIL)
    )
    return {
        "llm_base_url": config_value(
            "OPENAI_BASE_URL" if llm_provider != "Google Gemini" else "GEMINI_BASE_URL",
            DEFAULT_GEMINI_BASE_URL if llm_provider == "Google Gemini" else DEFAULT_LLM_BASE_URL,
        ),
        "llm_model": config_value(
            "OPENAI_MODEL" if llm_provider != "Google Gemini" else "GEMINI_MODEL",
            DEFAULT_GEMINI_MODEL if llm_provider == "Google Gemini" else DEFAULT_LLM_MODEL,
        ),
        "openalex_api_key": config_value("OPENALEX_API_KEY", DEFAULT_OPENALEX_API_KEY),
        "openalex_email": contact_email,
        "ncbi_api_key": config_value("NCBI_API_KEY"),
    }


def validate_llm_config(llm_api_key: str, llm_base_url: str, llm_model: str) -> bool:
    if not llm_api_key.strip():
        st.warning("请先填写 LLM API Key。")
        return False
    if not llm_base_url.strip():
        st.warning("请填写 LLM Base URL。")
        return False
    if not llm_model.strip():
        st.warning("请填写 LLM 模型名称。")
        return False
    return True


def render_markdown_download(label: str, content: str, file_name: str) -> None:
    st.download_button(
        label=label,
        data=content.encode("utf-8"),
        file_name=file_name,
        mime="text/markdown",
        use_container_width=True,
    )


def render_tool_result(title: str, content: str, file_name: str) -> None:
    st.markdown(f"### {title}")
    st.markdown(content)
    render_markdown_download("下载 Markdown", content, file_name)


def render_quick_start_guide() -> None:
    st.info(
        "使用前先在左侧填写 LLM API Key。"
        "然后选择下方功能：PDF精读把一篇论文变成结构化读书报告；引用追踪用于输入 DOI/OpenAlex ID 并导出 Excel；"
        "科研写作按任务拆分润色、审稿、引用数据、科研绘图和成果转化。"
    )
    with st.expander("快速使用说明", expanded=True):
        col1, col2, col3 = st.columns(3)
        with col1:
            st.markdown(
                "**1. 配置 Key**\n\n"
                "- 左侧选择 `OpenAI 兼容` 或 `Google Gemini`\n"
                "- 填写自己的 `LLM API Key`\n"
                "- 不确定时可点 `测试连接`"
            )
        with col2:
            st.markdown(
                "**2. 选择功能**\n\n"
                "- `PDF精读`：生成结构化读书报告\n"
                "- `引用追踪`：输入 DOI 或 Paper ID\n"
                "- `科研写作`：选择工作流后再提供材料"
            )
        with col3:
            st.markdown(
                "**3. 下载结果**\n\n"
                "- PDF 精读可下载 Word 报告\n"
                "- 引用追踪可下载 Excel\n"
                "- 云端不会保存你的上传文件"
            )


def apply_page_style() -> None:
    st.markdown(
        """
        <style>
        :root {
            --app-blue: var(--primary-color, #0f766e);
            --app-blue-dark: color-mix(in srgb, var(--primary-color, #0f766e) 82%, var(--text-color, #0f172a));
            --app-border: color-mix(in srgb, var(--text-color, #0f172a) 18%, transparent);
            --app-muted: color-mix(in srgb, var(--text-color, #0f172a) 68%, transparent);
            --app-surface: var(--secondary-background-color, #f0f2f6);
            --app-background: var(--background-color, #ffffff);
            --app-text: var(--text-color, #0f172a);
        }
        .block-container {
            max-width: 1080px;
            padding-top: 2rem;
            padding-bottom: 3rem;
        }
        .app-hero {
            padding: 1.15rem 1.3rem;
            margin-bottom: 1.25rem;
            border: 1px solid var(--app-border);
            border-radius: 16px;
            background: var(--app-background);
            background: linear-gradient(135deg, color-mix(in srgb, var(--app-blue) 10%, var(--app-background)) 0%, var(--app-background) 68%);
        }
        .app-kicker {
            color: var(--app-blue);
            font-size: .78rem;
            font-weight: 700;
            letter-spacing: .08em;
            text-transform: uppercase;
        }
        .app-hero h1 {
            font-size: clamp(1.75rem, 4vw, 2.5rem);
            margin: .3rem 0 .35rem 0;
            color: var(--app-text) !important;
        }
        .app-hero p { color: var(--app-muted) !important; margin: 0; }
        h1 {
            letter-spacing: 0;
            margin-bottom: 0.15rem;
            color: var(--app-text);
        }
        h2, h3 {
            letter-spacing: 0;
            color: var(--app-text);
        }
        div[data-testid="stCaptionContainer"] {
            color: var(--app-muted);
            line-height: 1.55;
        }
        div[data-testid="stTabs"] button {
            font-weight: 600;
            padding-top: 0.65rem;
            padding-bottom: 0.65rem;
        }
        div[data-testid="stTabs"] [aria-selected="true"] {
            color: var(--app-blue);
        }
        div[data-testid="stVerticalBlockBorderWrapper"] {
            border-color: var(--app-border);
            background: var(--app-background);
            border-radius: 12px;
        }
        div[data-testid="stAlert"] {
            border-radius: 10px;
        }
        div[data-testid="stFileUploader"] section {
            border-color: var(--app-border);
            background: var(--app-surface);
            border-radius: 10px;
        }
        .stDownloadButton button,
        .stButton button {
            border-radius: 9px;
            font-weight: 600;
            border-color: var(--app-border);
        }
        .stButton button[kind="primary"],
        .stDownloadButton button[kind="primary"] {
            background: var(--app-blue);
            border-color: var(--app-blue-dark);
        }
        div[data-testid="stDataFrame"] {
            border: 1px solid var(--app-border);
            border-radius: 6px;
            overflow: hidden;
        }
        section[data-testid="stSidebar"] .block-container {
            padding-top: 1.25rem;
        }
        section[data-testid="stSidebar"] {
            background: var(--app-surface);
            border-right: 1px solid var(--app-border);
        }
        section[data-testid="stSidebar"] h2,
        section[data-testid="stSidebar"] h3,
        section[data-testid="stSidebar"] div[role="radiogroup"] label p,
        section[data-testid="stSidebar"] div[data-testid="stCaptionContainer"] p {
            color: var(--app-text) !important;
        }
        section[data-testid="stSidebar"] div[role="radiogroup"] label {
            padding: .35rem .45rem;
            border-radius: 8px;
        }
        @media (max-width: 640px) {
            .block-container { padding-top: 1rem; }
            .app-hero { padding: 1rem; border-radius: 12px; }
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def request_json(
    url: str,
    *,
    params: Optional[Dict[str, Any]] = None,
    method: str = "GET",
    payload: Optional[Dict[str, Any]] = None,
    headers: Optional[Dict[str, str]] = None,
    timeout: int = 40,
    retries: int = 2,
) -> Dict[str, Any]:
    if params:
        query = urllib.parse.urlencode(params, doseq=True)
        url = f"{url}{'&' if '?' in url else '?'}{query}"

    request_headers = {
        "Accept": "application/json",
        "User-Agent": "streamlit-openalex-citation-analyzer/1.0",
    }
    if headers:
        request_headers.update(headers)

    body = None
    if payload is not None:
        body = json.dumps(payload, ensure_ascii=False).encode("utf-8")
        request_headers["Content-Type"] = "application/json"

    for attempt in range(retries + 1):
        request = urllib.request.Request(url, data=body, headers=request_headers, method=method)
        try:
            with urllib.request.urlopen(request, timeout=timeout) as response:
                raw = response.read().decode("utf-8", errors="replace")
                return json.loads(raw) if raw.strip() else {}
        except urllib.error.HTTPError as exc:
            raw = exc.read().decode("utf-8", errors="replace")
            if exc.code in {429, 500, 502, 503, 504} and attempt < retries:
                time.sleep(1.5 * (attempt + 1))
                continue
            raise ApiError(readable_http_error(exc.code, raw)) from exc
        except (urllib.error.URLError, TimeoutError, json.JSONDecodeError) as exc:
            if attempt < retries:
                time.sleep(1.5 * (attempt + 1))
                continue
            raise ApiError(str(exc)) from exc
    raise ApiError(f"请求失败：{url}")


def readable_http_error(status_code: int, raw_body: str) -> str:
    if status_code == 401:
        return (
            "LLM 服务认证失败（HTTP 401）。这通常表示 API Key 和 Base URL 不匹配，"
            "或 Key 已失效。请确认：OpenAI 官方 Key 使用 https://api.openai.com/v1；"
            "第三方/中转/学校平台 Key 必须填写对应平台提供的 Base URL；"
            "Gemini Key 请在侧边栏选择 Google Gemini 服务商。"
        )
    if status_code == 503:
        return (
            "LLM 服务暂不可用（HTTP 503）。这通常是模型当前高需求或服务临时拥堵，"
            "不是 PDF 或 Key 的问题。系统会自动重试并尝试备用模型；如果仍失败，"
            "请稍后再试或手动切换为 gemini-2.5-flash-lite / gemini-2.0-flash。"
        )
    if status_code == 429:
        return "LLM 服务限流（HTTP 429）。系统会自动重试；如果仍失败，请降低频率或稍后再试。"
    return f"HTTP {status_code}: {raw_body[:800]}"


def openalex_params(email: str, api_key: str, extra: Optional[Dict[str, Any]] = None) -> Dict[str, Any]:
    params: Dict[str, Any] = {}
    if email:
        params["mailto"] = email
    if api_key:
        params["api_key"] = api_key
    if extra:
        params.update(extra)
    return params


def restore_abstract(inverted_index: Any) -> str:
    if not isinstance(inverted_index, dict):
        return ""
    positioned: List[tuple[int, str]] = []
    for word, positions in inverted_index.items():
        if not isinstance(positions, list):
            continue
        for pos in positions:
            try:
                positioned.append((int(pos), str(word)))
            except (TypeError, ValueError):
                continue
    positioned.sort(key=lambda item: item[0])
    return strip_markup(" ".join(word for _, word in positioned))


def clean_doi(value: Any) -> str:
    if not value:
        return ""
    return normalize_doi(str(value))


def source_name(work: Dict[str, Any]) -> str:
    primary = work.get("primary_location") or {}
    source = primary.get("source") or {}
    return source.get("display_name") or ""


def open_access_pdf_url(work: Dict[str, Any]) -> str:
    locations = []
    primary = work.get("primary_location")
    if isinstance(primary, dict):
        locations.append(primary)
    locations.extend(work.get("locations") or [])
    best = work.get("best_oa_location")
    if isinstance(best, dict):
        locations.insert(0, best)

    for location in locations:
        if not isinstance(location, dict):
            continue
        pdf_url = location.get("pdf_url")
        if pdf_url:
            return str(pdf_url)
        landing = location.get("landing_page_url")
        if landing and str(landing).lower().endswith(".pdf"):
            return str(landing)
    return ""


def authors_text(work: Dict[str, Any], max_authors: int = 8) -> str:
    names = []
    for authorship in work.get("authorships") or []:
        author = authorship.get("author") or {}
        name = author.get("display_name")
        if name:
            names.append(str(name))
    suffix = " et al." if len(names) > max_authors else ""
    return "; ".join(names[:max_authors]) + suffix


def resolve_target_work(identifier: str, email: str, openalex_api_key: str) -> Dict[str, Any]:
    identifier = normalize_identifier(identifier)
    if is_openalex_id(identifier):
        path_id = clean_openalex_id(identifier)
        url = f"{OPENALEX_WORKS}/{path_id}"
    else:
        doi = normalize_doi(identifier)
        url = f"{OPENALEX_WORKS}/{urllib.parse.quote(f'https://doi.org/{doi}', safe=':/')}"
    return request_json(url, params=openalex_params(email, openalex_api_key), retries=2)


def fetch_citing_works(
    openalex_id: str,
    email: str,
    openalex_api_key: str,
    limit: int,
    per_page: int,
    progress,
    status_box,
) -> List[Dict[str, Any]]:
    works: List[Dict[str, Any]] = []
    cursor = "*"
    per_page = min(max(per_page, 1), 200)
    total_hint = None

    while len(works) < limit:
        page_size = min(per_page, limit - len(works))
        params = openalex_params(
            email,
            openalex_api_key,
            {
                "filter": f"cites:{openalex_id}",
                "per-page": page_size,
                "cursor": cursor,
            },
        )
        page = request_json(OPENALEX_WORKS, params=params, retries=2)
        results = page.get("results") or []
        meta = page.get("meta") or {}
        total_hint = total_hint or meta.get("count")
        works.extend(results)

        progress.progress(min(len(works) / max(limit, 1), 0.35), text=f"正在拉取文献网络... 已获取 {len(works)} 篇")
        status_box.write(f"已获取 {len(works)} 篇引用文献，OpenAlex 估计总数：{total_hint or '未知'}")

        next_cursor = meta.get("next_cursor")
        if not results or not next_cursor:
            break
        cursor = next_cursor
        time.sleep(0.1)
    return works


def work_to_record(work: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "title": strip_markup(work.get("title") or work.get("display_name") or ""),
        "year": work.get("publication_year") or "",
        "doi": clean_doi(work.get("doi")),
        "abstract": restore_abstract(work.get("abstract_inverted_index")),
        "openalex_id": clean_openalex_id(str(work.get("id") or "")),
        "cited_by_count": work.get("cited_by_count") or 0,
        "source": source_name(work),
        "authors": authors_text(work),
        "oa_pdf_url": open_access_pdf_url(work),
        "pmid": "",
        "mesh": "",
        "evidence_level": "元数据",
        "citation_context": "",
        "context_source": "",
        "parse_status": "OpenAlex 元数据",
    }


def build_output_row(record: Dict[str, Any]) -> Dict[str, Any]:
    return {
        "文献标题": record.get("title", ""),
        "年份": record.get("year", ""),
        "DOI": record.get("doi", ""),
        "摘要": record.get("abstract", ""),
        "OpenAlex ID": record.get("openalex_id", ""),
        "引用数": record.get("cited_by_count", 0),
        "来源期刊": record.get("source", ""),
        "作者": record.get("authors", ""),
        "证据等级": record.get("evidence_level", ""),
        "引用上下文": record.get("citation_context", ""),
        "上下文来源": record.get("context_source", ""),
        "PMID": record.get("pmid", ""),
        "MeSH": record.get("mesh", ""),
        "开放全文URL": record.get("oa_pdf_url", ""),
        "解析状态": record.get("parse_status", ""),
    }


def dataframe_to_excel_bytes(df: pd.DataFrame) -> bytes:
    buffer = BytesIO()
    with pd.ExcelWriter(buffer, engine="openpyxl") as writer:
        df.to_excel(writer, index=False, sheet_name="Citation_Analysis")
        doi_columns = [column for column in ["文献标题", "年份", "DOI", "PMID", "OpenAlex ID", "证据等级"] if column in df.columns]
        df[doi_columns].to_excel(writer, index=False, sheet_name="DOI_List")
    buffer.seek(0)
    return buffer.getvalue()


def fetch_crossref_metadata(record: Dict[str, Any], email: str = "") -> Dict[str, str]:
    doi = record.get("doi") or ""
    title = record.get("title") or ""
    if not doi and not title:
        return {}

    try:
        if doi:
            url = f"https://api.crossref.org/works/{urllib.parse.quote(doi, safe='')}"
            response = request_json(url, params={"mailto": email} if email else None, retries=1)
            item = response.get("message") or {}
        else:
            response = request_json(
                "https://api.crossref.org/works",
                params={"query.title": title, "rows": 1, **({"mailto": email} if email else {})},
                retries=1,
            )
            items = ((response.get("message") or {}).get("items") or [])
            item = items[0] if items else {}
    except Exception:
        return {}

    if not item:
        return {}
    container = item.get("container-title") or []
    return {
        "doi": clean_doi(item.get("DOI")),
        "source": container[0] if container else "",
        "year": (((item.get("published-print") or item.get("published-online") or {}).get("date-parts") or [[None]])[0][0] or ""),
    }


def fetch_pubmed_metadata_by_doi(doi: str, email: str = "") -> Dict[str, str]:
    if not doi:
        return {}
    params = {
        "db": "pubmed",
        "term": f"{doi}[DOI]",
        "retmode": "json",
        "tool": "biomed-literature-workbench",
    }
    if email:
        params["email"] = email
    try:
        search = request_json("https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esearch.fcgi", params=params, retries=1)
        ids = (((search.get("esearchresult") or {}).get("idlist")) or [])
        if not ids:
            return {}
        pmid = str(ids[0])
        summary = request_json(
            "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/esummary.fcgi",
            params={"db": "pubmed", "id": pmid, "retmode": "json", "tool": "biomed-literature-workbench", **({"email": email} if email else {})},
            retries=1,
        )
        item = ((summary.get("result") or {}).get(pmid) or {})
        mesh_terms = []
        for term in item.get("meshheadinglist") or []:
            if isinstance(term, str):
                mesh_terms.append(term)
            elif isinstance(term, dict):
                mesh_terms.append(str(term.get("name") or term.get("term") or ""))
        return {"pmid": pmid, "mesh": "; ".join(term for term in mesh_terms if term)}
    except Exception:
        return {}


def post_pdf_to_grobid(pdf_bytes: bytes) -> str:
    base_url = os.getenv("GROBID_BASE_URL", "").rstrip("/")
    if not base_url:
        raise ApiError("未配置 GROBID_BASE_URL。")

    boundary = f"----biomed-workbench-{int(time.time() * 1000)}"
    body = b"".join(
        [
            f"--{boundary}\r\n".encode("utf-8"),
            b'Content-Disposition: form-data; name="input"; filename="paper.pdf"\r\n',
            b"Content-Type: application/pdf\r\n\r\n",
            pdf_bytes,
            b"\r\n",
            f"--{boundary}--\r\n".encode("utf-8"),
        ]
    )
    request = urllib.request.Request(
        f"{base_url}/api/processFulltextDocument",
        data=body,
        headers={
            "Content-Type": f"multipart/form-data; boundary={boundary}",
            "Accept": "application/xml",
            "User-Agent": "biomed-literature-workbench/1.0",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=45) as response:
            return response.read().decode("utf-8", errors="replace")
    except Exception as exc:
        raise ApiError(f"GROBID 解析失败：{exc}") from exc


def tei_to_plain_text(tei_xml: str) -> str:
    text = re.sub(r"<ref\b[^>]*>", " ", tei_xml or "", flags=re.I)
    text = re.sub(r"</ref>", " ", text, flags=re.I)
    text = re.sub(r"<[^>]+>", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def enrich_citing_work(record: Dict[str, Any], email: str = "") -> Dict[str, Any]:
    enriched = dict(record)
    status = [enriched.get("parse_status") or "OpenAlex 元数据"]

    crossref = {}
    if not enriched.get("doi") or not enriched.get("source") or not enriched.get("year"):
        crossref = fetch_crossref_metadata(enriched, email)
    if crossref:
        if not enriched.get("doi") and crossref.get("doi"):
            enriched["doi"] = crossref["doi"]
        if not enriched.get("source") and crossref.get("source"):
            enriched["source"] = crossref["source"]
        if not enriched.get("year") and crossref.get("year"):
            enriched["year"] = crossref["year"]
        status.append("Crossref 补全")

    pubmed = fetch_pubmed_metadata_by_doi(enriched.get("doi") or "", email) if enriched.get("doi") else {}
    if pubmed:
        enriched["pmid"] = pubmed.get("pmid", "")
        enriched["mesh"] = pubmed.get("mesh", "")
        status.append("PubMed 补全")

    enriched["parse_status"] = "；".join(status)
    return enriched


def download_open_pdf_bytes(pdf_url: str, max_bytes: int = 8_000_000) -> bytes:
    if not pdf_url:
        raise ApiError("无开放全文 PDF URL。")
    request = urllib.request.Request(pdf_url, headers={"User-Agent": "biomed-literature-workbench/1.0"})
    try:
        with urllib.request.urlopen(request, timeout=25) as response:
            content_type = response.headers.get("Content-Type", "")
            raw = response.read(max_bytes + 1)
        if len(raw) > max_bytes:
            raise ApiError("开放全文 PDF 过大，已跳过。")
        if "pdf" not in content_type.lower() and not pdf_url.lower().endswith(".pdf"):
            raise ApiError("开放全文链接不是 PDF，已跳过。")
    except Exception as exc:
        if isinstance(exc, ApiError):
            raise
        raise ApiError(f"开放全文下载失败：{exc}") from exc
    return raw


def download_open_pdf_text(pdf_url: str, max_bytes: int = 8_000_000) -> str:
    raw = download_open_pdf_bytes(pdf_url, max_bytes)

    try:
        return pypdf_extract_full_text(raw, max_pages=40)
    except Exception as exc:
        raise ApiError(f"开放全文文本解析失败：{exc}") from exc


def extract_context_window(text: str, needle: str, window: int = 900) -> str:
    if not text or not needle:
        return ""
    index = text.lower().find(needle.lower())
    if index < 0:
        return ""
    start = max(0, index - window)
    end = min(len(text), index + len(needle) + window)
    return normalize_pdf_text(text[start:end], 2200)


def extract_citation_context(target_doi: str, citing_record: Dict[str, Any]) -> Dict[str, str]:
    pdf_url = citing_record.get("oa_pdf_url") or ""
    if not pdf_url:
        return {
            "evidence_level": "摘要/元数据",
            "citation_context": "",
            "context_source": "无开放全文",
            "parse_status": "无开放全文 PDF URL",
        }

    try:
        pdf_bytes = download_open_pdf_bytes(pdf_url)
        parser_source = "pypdf"
        try:
            tei_xml = post_pdf_to_grobid(pdf_bytes)
            full_text = tei_to_plain_text(tei_xml)
            parser_source = "GROBID TEI"
            if not full_text:
                raise ApiError("GROBID 未返回可读文本。")
        except Exception:
            full_text = pypdf_extract_full_text(pdf_bytes, max_pages=40)
    except Exception as exc:
        return {
            "evidence_level": "摘要/元数据",
            "citation_context": "",
            "context_source": "开放全文解析失败",
            "parse_status": str(exc),
        }

    context = extract_context_window(full_text, normalize_doi(target_doi))
    source = "开放全文 DOI 匹配"
    if not context:
        context = extract_context_window(full_text, target_doi)
    if not context:
        source = f"开放全文可读但未定位目标 DOI（{parser_source}）"
    else:
        source = f"{source}（{parser_source}）"

    return {
        "evidence_level": "全文上下文" if context else "开放全文未定位",
        "citation_context": context,
        "context_source": source,
        "parse_status": "开放全文 PDF 解析成功" if context else "开放全文解析成功，但未匹配目标 DOI",
    }


def selected_pdf_page_indexes(page_count: int) -> List[int]:
    """Use the same high-signal page window for text and image reading."""
    if page_count <= 0:
        raise ApiError("PDF 中没有可读取页面。")
    selected = list(range(min(3, page_count)))
    tail_start = max(0, page_count - 2)
    selected.extend(range(tail_start, page_count))
    return sorted(set(selected))


def normalize_pdf_text(text: str, max_chars: int = 24000) -> str:
    return compact_text(re.sub(r"\s+", " ", text or "").strip(), max_chars)


def extract_sections_from_text(text: str) -> Dict[str, str]:
    section_patterns = {
        "title_abstract": r"(?is)^(.{0,5000}?)(?=\bintroduction\b|\bbackground\b|\bmaterials and methods\b|\bmethods\b|\bresults\b)",
        "introduction": r"(?is)\b(?:introduction|background)\b(.{0,6000}?)(?=\bmaterials and methods\b|\bmethods\b|\bresults\b|\bdiscussion\b)",
        "methods": r"(?is)\b(?:materials and methods|methods|methodology|experimental procedures)\b(.{0,7000}?)(?=\bresults\b|\bdiscussion\b|\bconclusion\b|\breferences\b)",
        "results": r"(?is)\bresults\b(.{0,8000}?)(?=\bdiscussion\b|\bconclusion\b|\bmaterials and methods\b|\bmethods\b|\breferences\b)",
        "discussion_conclusion": r"(?is)\b(?:discussion|conclusion|conclusions)\b(.{0,7000}?)(?=\breferences\b|\backnowledg)",
        "references": r"(?is)\breferences\b(.{0,5000})$",
    }
    sections: Dict[str, str] = {}
    for name, pattern in section_patterns.items():
        match = re.search(pattern, text or "")
        if match:
            sections[name] = normalize_pdf_text(match.group(0), 5000)
    return sections


FIGURE_TABLE_MARKER_RE = re.compile(
    r"(?im)(?:^|\n)\s*("
    r"(?:Extended\s+Data\s+Fig(?:ure)?|Supplementary\s+Fig(?:ure)?|Supplementary\s+Table|"
    r"Fig(?:ure)?|Table|Scheme)"
    r"\.?\s*[A-Za-z]?\d+[A-Za-z]?(?:[-–]\d+[A-Za-z]?)?"
    r")\s*[:.|：\-–]?\s*"
)


def normalize_figure_id(value: Any) -> str:
    text = re.sub(r"\s+", " ", str(value or "")).strip().rstrip(".:：|-–")
    text = re.sub(r"(?i)^figure\b", "Fig.", text)
    text = re.sub(r"(?i)^fig\b(?!\.)", "Fig.", text)
    text = re.sub(r"(?i)^table\b", "Table", text)
    text = re.sub(r"(?i)^scheme\b", "Scheme", text)
    return text


def figure_id_key(value: Any) -> str:
    text = normalize_figure_id(value).lower()
    text = text.replace("figure", "fig")
    return re.sub(r"[^a-z0-9]+", "", text)


def extract_page_number_near(text: str, position: int) -> str:
    prefix = text[:position]
    matches = list(re.finditer(r"\[Page\s+(\d+)\]", prefix, flags=re.I))
    return matches[-1].group(1) if matches else ""


def trim_caption_segment(segment: str, max_chars: int = 1600) -> str:
    segment = re.sub(r"\[Page\s+\d+\]", " ", segment, flags=re.I)
    segment = re.sub(r"\s+", " ", segment).strip()
    segment = re.sub(
        r"(?i)\s+(references|acknowledg(?:e)?ments|author contributions|competing interests|methods)\s+.*$",
        "",
        segment,
    ).strip()
    return compact_text(segment, max_chars)


def extract_figure_table_legends(text: str, max_items: int = 80) -> List[Dict[str, str]]:
    raw_text = text or ""
    matches = list(FIGURE_TABLE_MARKER_RE.finditer(raw_text))
    legends: List[Dict[str, str]] = []
    seen: set[str] = set()

    for index, match in enumerate(matches):
        figure_id = normalize_figure_id(match.group(1))
        key = figure_id_key(figure_id)
        if not key or key in seen:
            continue

        next_start = matches[index + 1].start(1) if index + 1 < len(matches) else len(raw_text)
        segment = raw_text[match.start(1) : min(next_start, match.start(1) + 2200)]
        caption = trim_caption_segment(segment)
        if len(caption) < len(figure_id) + 20:
            continue

        seen.add(key)
        legends.append(
            {
                "figure_id": figure_id,
                "caption": caption,
                "page": extract_page_number_near(raw_text, match.start(1)),
            }
        )
        if len(legends) >= max_items:
            break
    return legends


def extract_figure_table_clues(text: str, max_items: int = 18) -> List[str]:
    legends = extract_figure_table_legends(text, max_items=max_items)
    if legends:
        return [item["caption"] for item in legends]

    clues: List[str] = []
    for match in re.finditer(r"(?is)\b(?:fig(?:ure)?\.?\s*\d+[a-z]?|table\s*\d+[a-z]?)\b.{0,700}", text or ""):
        clue = normalize_pdf_text(match.group(0), 900)
        if clue and clue not in clues:
            clues.append(clue)
        if len(clues) >= max_items:
            break
    return clues


def pypdf_extract_full_text(pdf_bytes: bytes, max_pages: Optional[int] = None) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise ApiError("当前环境缺少 pypdf，请先安装：pip install pypdf") from exc

    try:
        reader = PdfReader(BytesIO(pdf_bytes))
        page_texts: List[str] = []
        pages = reader.pages if max_pages is None else reader.pages[:max_pages]
        for page_index, page in enumerate(pages, start=1):
            text = page.extract_text() or ""
            if text.strip():
                page_texts.append(f"[Page {page_index}]\n{text}")
        cleaned = "\n\n".join(page_texts).strip()
    except Exception as exc:
        raise ApiError(f"pypdf 全文解析失败：{exc}") from exc

    if not cleaned:
        raise ApiError("pypdf 未能从 PDF 中提取到文本。")
    return cleaned


def extract_pdf_with_pymupdf4llm(pdf_bytes: bytes) -> Dict[str, Any]:
    try:
        import pymupdf4llm
    except ImportError as exc:
        raise ApiError("PyMuPDF4LLM 未安装，已降级为 pypdf 结构抽取。") from exc

    temp_path = None
    try:
        import tempfile

        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            temp_file.write(pdf_bytes)
            temp_path = temp_file.name
        markdown = pymupdf4llm.to_markdown(temp_path)
        if isinstance(markdown, list):
            markdown = "\n\n".join(str(item) for item in markdown)
        markdown = str(markdown or "").strip()
        if not markdown:
            raise ApiError("PyMuPDF4LLM 未返回有效 Markdown。")
        sections = extract_sections_from_text(markdown)
        figure_table_legends = extract_figure_table_legends(markdown)
        return {
            "mode": PDF_PARSE_ENHANCED,
            "parser": "PyMuPDF4LLM",
            "markdown": compact_text(markdown, 28000),
            "sections": sections,
            "doi": extract_doi_from_text(markdown),
            "figure_table_legends": figure_table_legends,
            "figure_table_clues": [item["caption"] for item in figure_table_legends[:18]]
            or extract_figure_table_clues(markdown),
            "evidence_level": "增强解析",
            "status": "PyMuPDF4LLM Markdown 解析成功",
        }
    finally:
        if temp_path:
            try:
                os.unlink(temp_path)
            except OSError:
                pass


def extract_pdf_with_docling(pdf_bytes: bytes) -> Dict[str, Any]:
    try:
        from docling.document_converter import DocumentConverter
    except ImportError as exc:
        raise ApiError("Docling 未安装，已降级为启发式结构抽取。") from exc

    temp_path = None
    try:
        import tempfile

        with tempfile.NamedTemporaryFile(delete=False, suffix=".pdf") as temp_file:
            temp_file.write(pdf_bytes)
            temp_path = temp_file.name
        result = DocumentConverter().convert(temp_path)
        document = result.document
        markdown = document.export_to_markdown()
        sections = extract_sections_from_text(markdown)
        figure_table_legends = extract_figure_table_legends(markdown)
        return {
            "mode": PDF_PARSE_STRUCTURED,
            "parser": "Docling",
            "markdown": compact_text(markdown, 32000),
            "sections": sections,
            "doi": extract_doi_from_text(markdown),
            "figure_table_legends": figure_table_legends,
            "figure_table_clues": [item["caption"] for item in figure_table_legends[:18]]
            or extract_figure_table_clues(markdown),
            "evidence_level": "结构化解析",
            "status": "Docling 结构化解析成功",
        }
    finally:
        if temp_path:
            try:
                os.unlink(temp_path)
            except OSError:
                pass


def build_pdf_content_from_text(text: str, mode: str, parser: str, status: str) -> Dict[str, Any]:
    sections = extract_sections_from_text(text)
    figure_table_legends = extract_figure_table_legends(text)
    return {
        "mode": mode,
        "parser": parser,
        "markdown": compact_text(text, 28000),
        "sections": sections,
        "doi": extract_doi_from_text(text),
        "figure_table_legends": figure_table_legends,
        "figure_table_clues": [item["caption"] for item in figure_table_legends[:18]] or extract_figure_table_clues(text),
        "evidence_level": "全文结构推断" if sections else "核心页文本",
        "status": status,
    }


def extract_pdf_content(pdf_bytes: bytes, mode: str) -> Dict[str, Any]:
    if mode == PDF_PARSE_FAST:
        text = extract_core_pdf_text(pdf_bytes)
        return build_pdf_content_from_text(text, mode, "pdfplumber", "快速模式：核心页文本解析成功")

    if mode == PDF_PARSE_ENHANCED:
        try:
            return extract_pdf_with_pymupdf4llm(pdf_bytes)
        except ApiError as exc:
            fallback_text = pypdf_extract_full_text(pdf_bytes)
            content = build_pdf_content_from_text(
                fallback_text,
                mode,
                "pypdf fallback",
                f"增强模式降级：{exc}",
            )
            content["evidence_level"] = "增强降级"
            return content

    if mode == PDF_PARSE_STRUCTURED:
        try:
            return extract_pdf_with_docling(pdf_bytes)
        except ApiError as exc:
            fallback_text = pypdf_extract_full_text(pdf_bytes)
            content = build_pdf_content_from_text(
                fallback_text,
                mode,
                "heuristic sections",
                f"结构化模式降级：{exc}",
            )
            content["evidence_level"] = "结构化降级"
            return content

    raise ApiError(f"未知 PDF 解析模式：{mode}")


def pdf_content_to_prompt_text(content: Dict[str, Any], max_chars: int = 18000) -> str:
    sections = content.get("sections") or {}
    parts = [
        f"解析器：{content.get('parser', '')}",
        f"证据等级：{content.get('evidence_level', '')}",
        f"解析状态：{content.get('status', '')}",
    ]
    legends = content.get("figure_table_legends") or []
    if legends:
        parts.append("\n全文图例/图注：")
        for item in legends[:30]:
            page = f" Page {item.get('page')}" if item.get("page") else ""
            parts.append(f"- {item.get('figure_id', '')}{page}: {item.get('caption', '')}")

    if sections:
        parts.append("结构化片段：")
        for name, section_text in sections.items():
            parts.append(f"\n## {name}\n{section_text}")
    else:
        parts.append("正文片段：")
        parts.append(str(content.get("markdown") or ""))

    clues = content.get("figure_table_clues") or []
    if clues:
        parts.append("\n图表/表格线索：")
        for clue in clues[:12]:
            parts.append(f"- {clue}")
    return compact_text("\n".join(parts), max_chars)


def run_analysis(
    identifier: str,
    openalex_api_key: str,
    email: str,
    max_papers: int,
    citation_mode: str,
    progress,
    title_log,
) -> pd.DataFrame:
    progress.progress(0.03, text="正在解析目标文献...")
    target = resolve_target_work(identifier, email, openalex_api_key)
    openalex_id = clean_openalex_id(target.get("id") or "")
    target_doi = clean_doi(target.get("doi")) or normalize_doi(identifier)
    if not openalex_id:
        raise ApiError("OpenAlex 未能解析目标文献 ID。")

    progress.progress(0.1, text="正在拉取文献网络...")
    citing_works = fetch_citing_works(openalex_id, email, openalex_api_key, max_papers, 200, progress, title_log)

    rows: List[Dict[str, Any]] = []
    total = len(citing_works)
    for idx, work in enumerate(citing_works, start=1):
        record = work_to_record(work)
        title_log.write(f"正在整理第 {idx}/{total} 篇文献：{record['title']}")
        progress.progress(0.35 + 0.6 * idx / max(total, 1), text=f"正在整理第 {idx}/{total} 篇文献...")
        record = enrich_citing_work(record, email)
        if citation_mode == CITATION_CONTEXT_MODE:
            title_log.write(f"正在尝试开放全文上下文：{record['title']}")
            context = extract_citation_context(target_doi, record)
            record.update(context)
        rows.append(build_output_row(record))

    df = pd.DataFrame(rows, columns=OUTPUT_COLUMNS)
    if not df.empty:
        df = df.sort_values(["年份", "引用数"], ascending=[False, False]).reset_index(drop=True)
    progress.progress(1.0, text="分析完成")
    return df


def extract_core_pdf_text(pdf_bytes: bytes) -> str:
    """Extract only the high-signal pages: first 3 pages and last 2 pages."""
    try:
        import pdfplumber
    except ImportError as exc:
        raise ApiError("当前环境缺少 pdfplumber，请先安装：pip install pdfplumber") from exc

    try:
        with pdfplumber.open(BytesIO(pdf_bytes)) as pdf:
            page_count = len(pdf.pages)
            selected = selected_pdf_page_indexes(page_count)

            page_texts: List[str] = []
            for page_index in selected:
                text = pdf.pages[page_index].extract_text() or ""
                if text.strip():
                    page_texts.append(f"[Page {page_index + 1}]\n{text}")
    except Exception as exc:
        if isinstance(exc, ApiError):
            raise
        raise ApiError(f"PDF 解析失败，请确认文件未损坏且不是纯扫描图片：{exc}") from exc

    cleaned = re.sub(r"\s+", " ", "\n".join(page_texts)).strip()
    if not cleaned:
        raise ApiError("未能从 PDF 中提取到文本。若这是扫描版 PDF，请先 OCR。")
    return compact_text(cleaned, 12000)


def compress_pil_image_to_jpeg(image: Any) -> bytes:
    """Compress a rendered page while keeping it readable for vision models."""
    try:
        from PIL import Image
    except ImportError as exc:
        raise ApiError("当前环境缺少 Pillow，请先安装：pip install pillow") from exc

    if image.mode not in {"RGB", "L"}:
        image = image.convert("RGB")
    elif image.mode == "L":
        image = image.convert("RGB")

    quality = PDF_IMAGE_JPEG_QUALITY
    while True:
        buffer = BytesIO()
        image.save(buffer, format="JPEG", quality=quality, optimize=True)
        data = buffer.getvalue()
        if len(data) <= PDF_IMAGE_MAX_BYTES or quality <= 45:
            return data
        quality -= 8


def render_core_pdf_pages_as_images(pdf_bytes: bytes) -> List[Dict[str, Any]]:
    """Render first 3 and last 2 PDF pages into compact JPEG images."""
    try:
        import pypdfium2 as pdfium
    except ImportError as exc:
        raise ApiError("当前环境缺少 pypdfium2，请先安装：pip install pypdfium2 pillow") from exc

    try:
        pdf = pdfium.PdfDocument(pdf_bytes)
        selected = selected_pdf_page_indexes(len(pdf))
        images: List[Dict[str, Any]] = []

        for page_index in selected:
            page = pdf[page_index]
            bitmap = page.render(scale=PDF_IMAGE_RENDER_SCALE)
            pil_image = bitmap.to_pil()
            image_bytes = compress_pil_image_to_jpeg(pil_image)
            images.append(
                {
                    "page": page_index + 1,
                    "mime_type": "image/jpeg",
                    "data": image_bytes,
                }
            )
        return images
    except Exception as exc:
        if isinstance(exc, ApiError):
            raise
        raise ApiError(f"PDF 图片渲染失败，请确认文件未损坏：{exc}") from exc


def call_llm(
    llm_api_key: str,
    llm_provider: str,
    llm_base_url: str,
    llm_model: str,
    system_prompt: str,
    user_prompt: str,
    timeout: int = 120,
) -> str:
    if llm_provider == "Google Gemini":
        return call_gemini_llm(
            llm_api_key=llm_api_key,
            llm_base_url=llm_base_url,
            llm_model=llm_model,
            system_prompt=system_prompt,
            user_prompt=user_prompt,
            timeout=timeout,
        )
    return call_openai_compatible_llm(
        llm_api_key=llm_api_key,
        llm_base_url=llm_base_url,
        llm_model=llm_model,
        system_prompt=system_prompt,
        user_prompt=user_prompt,
        timeout=timeout,
    )


@st.cache_data(ttl=21600, max_entries=256, show_spinner=False)
def cached_literature_search(
    query: str,
    selected_sources: tuple[str, ...],
    start_year: int,
    end_year: int,
    limit: int,
    smart_rewrite: bool,
    llm_provider: str,
    llm_base_url: str,
    llm_model: str,
    cache_scope: str,
    _credentials: Dict[str, str],
    _llm_api_key: str,
) -> FederatedSearchResult:
    """Cache one federated search without hashing or storing credentials in its key."""
    del cache_scope
    llm_callback = None
    if smart_rewrite and _llm_api_key.strip():
        def llm_callback(prompt: str) -> str:
            return call_llm(
                llm_api_key=_llm_api_key.strip(),
                llm_provider=llm_provider,
                llm_base_url=llm_base_url.strip(),
                llm_model=llm_model.strip(),
                system_prompt=(
                    "你是生物医学信息检索专家。只返回一个 JSON 对象，键必须严格为 "
                    "pubmed、europe_pmc、openalex、crossref、terms_en、terms_zh；"
                    "不得生成论文、DOI、PMID 或用户未提供的日期限制。"
                ),
                user_prompt=prompt,
                timeout=90,
            )

    return search_literature(
        query=query,
        selected_sources=selected_sources,
        start_year=start_year,
        end_year=end_year,
        limit=limit,
        credentials=_credentials,
        smart_rewrite=smart_rewrite,
        llm_callback=llm_callback,
    )


def call_openai_compatible_llm(
    llm_api_key: str,
    llm_base_url: str,
    llm_model: str,
    system_prompt: str,
    user_prompt: str,
    timeout: int = 120,
) -> str:
    payload = {
        "model": llm_model,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        "temperature": 0.2,
    }
    base_url = llm_base_url.rstrip("/")
    response = request_json(
        f"{base_url}/chat/completions",
        method="POST",
        payload=payload,
        headers={"Authorization": f"Bearer {llm_api_key}"},
        timeout=timeout,
        retries=1,
    )
    return (((response.get("choices") or [{}])[0].get("message") or {}).get("content") or "").strip()


def call_gemini_llm(
    llm_api_key: str,
    llm_base_url: str,
    llm_model: str,
    system_prompt: str,
    user_prompt: str,
    timeout: int = 120,
) -> str:
    payload = {
        "systemInstruction": {
            "parts": [{"text": system_prompt}],
        },
        "contents": [
            {
                "role": "user",
                "parts": [{"text": user_prompt}],
            }
        ],
        "generationConfig": {
            "temperature": 0.2,
        },
    }
    return execute_gemini_payload(llm_api_key, llm_base_url, llm_model, payload, timeout)


def call_gemini_llm_with_images(
    llm_api_key: str,
    llm_base_url: str,
    llm_model: str,
    system_prompt: str,
    user_prompt: str,
    images: List[Dict[str, Any]],
    timeout: int = 180,
) -> str:
    if not images:
        raise ApiError("图片模式未生成可发送给 Gemini 的页面图片。")

    parts: List[Dict[str, Any]] = [{"text": user_prompt}]
    for image in images:
        parts.append({"text": f"下面是 PDF 第 {image['page']} 页的页面截图："})
        parts.append(
            {
                "inline_data": {
                    "mime_type": image["mime_type"],
                    "data": base64.b64encode(image["data"]).decode("ascii"),
                }
            }
        )

    payload = {
        "systemInstruction": {
            "parts": [{"text": system_prompt}],
        },
        "contents": [
            {
                "role": "user",
                "parts": parts,
            }
        ],
        "generationConfig": {
            "temperature": 0.2,
        },
    }
    return execute_gemini_payload(llm_api_key, llm_base_url, llm_model, payload, timeout)


def execute_gemini_payload(
    llm_api_key: str,
    llm_base_url: str,
    llm_model: str,
    payload: Dict[str, Any],
    timeout: int,
) -> str:
    base_url = llm_base_url.rstrip("/")
    model_candidates = gemini_model_candidates(llm_model)
    last_error = ""
    for model in model_candidates:
        endpoint = f"{base_url}/models/{urllib.parse.quote(model, safe='')}:generateContent"
        for attempt in range(4):
            try:
                response = request_json(
                    endpoint,
                    params={"key": llm_api_key},
                    method="POST",
                    payload=payload,
                    timeout=timeout,
                    retries=0,
                )
                candidates = response.get("candidates") or []
                if not candidates:
                    raise ApiError(f"Gemini 未返回候选结果：{str(response)[:500]}")
                parts = (((candidates[0].get("content") or {}).get("parts")) or [])
                text = "\n".join(str(part.get("text", "")) for part in parts if isinstance(part, dict))
                if text.strip():
                    return text.strip()
                raise ApiError(f"Gemini 返回了空文本：{str(response)[:500]}")
            except ApiError as exc:
                last_error = str(exc)
                if should_retry_llm_error(last_error) and attempt < 3:
                    time.sleep(min(20, 2 ** attempt * 2))
                    continue
                break
    raise ApiError(f"Gemini 多次重试和备用模型切换后仍失败：{last_error}")


def gemini_model_candidates(selected_model: str) -> List[str]:
    models = [selected_model.strip()] if selected_model.strip() else []
    for model in GEMINI_FALLBACK_MODELS:
        if model not in models:
            models.append(model)
    return models


def should_retry_llm_error(message: str) -> bool:
    retry_markers = ["HTTP 503", "HTTP 429", "暂不可用", "限流", "UNAVAILABLE", "high demand"]
    return any(marker in message for marker in retry_markers)


def analyze_pdf_deep_reading(
    pdf_text: str,
    llm_api_key: str,
    llm_provider: str,
    llm_base_url: str,
    llm_model: str,
) -> str:
    prompt = f"""
你是一个资深的生物医药方向研究员。请阅读以下提供的学术论文文本片段，并严格按照以下结构输出中文报告：

【主要内容】：用一段话总结文章做了什么。

【核心创新点】：列出 1-3 条该文献区别于过往研究的创新（如新的组装技术、新的酶突变位点等）。

【核心数据与实验结果】：提取支撑其结论的最重要的数据。

【期刊评估 (IF)】：请根据文本中出现的期刊名称，评估并给出该期刊近期的近似影响因子 (Impact Factor)。如果文本中找不到期刊名，请根据文章水平给出一个合理的期刊等级预测（如一区/二区）。

论文文本片段如下：
{pdf_text}
""".strip()
    return call_llm(
        llm_api_key=llm_api_key,
        llm_provider=llm_provider,
        llm_base_url=llm_base_url,
        llm_model=llm_model,
        system_prompt="你是一个严谨、克制、重视证据边界的生物医药方向研究员。",
        user_prompt=prompt,
    )


def analyze_pdf_content_deep_reading(
    content: Dict[str, Any],
    llm_api_key: str,
    llm_provider: str,
    llm_base_url: str,
    llm_model: str,
) -> str:
    prompt_text = pdf_content_to_prompt_text(content)
    prompt = f"""
你是一个资深的生物医药方向研究员。请阅读以下结构化论文内容，并严格按照以下结构输出中文报告：

【主要内容】：用一段话总结文章做了什么。

【核心创新点】：列出 1-3 条该文献区别于过往研究的创新（如新的组装技术、新的酶突变位点等）。

【核心数据与实验结果】：提取支撑其结论的最重要数据；请优先使用 results、figure/table clues 和 methods 中的证据。

【实验方法与技术路线】：尽量提取关键实验方法、测序/富集/扩增/酶工程/组装策略、样本和分析流程。

【期刊评估 (IF)】：根据文本中的期刊名称给出近期近似影响因子；找不到期刊名时给出合理等级预测并说明不确定性。

论文结构化内容如下：
{prompt_text}
""".strip()
    return call_llm(
        llm_api_key=llm_api_key,
        llm_provider=llm_provider,
        llm_base_url=llm_base_url,
        llm_model=llm_model,
        system_prompt="你是一个严谨、克制、重视证据边界的生物医药方向研究员。",
        user_prompt=prompt,
    )


def analyze_pdf_deep_reading_with_images(
    images: List[Dict[str, Any]],
    llm_api_key: str,
    llm_provider: str,
    llm_base_url: str,
    llm_model: str,
) -> str:
    if llm_provider != "Google Gemini":
        raise ApiError("当前图片识别模式请切换到 Google Gemini。")

    page_list = "、".join(str(image["page"]) for image in images)
    prompt = f"""
你是一个资深的生物医药方向研究员。请直接阅读下面提供的 PDF 页面截图，页面范围为第 {page_list} 页。

这些截图来自论文的高信号页面（前 3 页和最后 2 页）。请尽量从标题、摘要、图表、结论、方法和参考信息中提取证据，并严格按照以下结构输出中文报告：

【主要内容】：用一段话总结文章做了什么。

【核心创新点】：列出 1-3 条该文献区别于过往研究的创新（如新的组装技术、新的酶突变位点等）。

【核心数据与实验结果】：提取支撑其结论的最重要的数据；如果图表中有关键数值、样本量、效率、准确率或酶学指标，请优先列出。

【期刊评估 (IF)】：请根据截图中出现的期刊名称，评估并给出该期刊近期的近似影响因子 (Impact Factor)。如果截图中找不到期刊名，请根据文章水平给出一个合理的期刊等级预测（如一区/二区）。

请忽略页眉页脚、水印、版权声明等非生物学相关冗余信息。
""".strip()
    return call_gemini_llm_with_images(
        llm_api_key=llm_api_key,
        llm_base_url=llm_base_url,
        llm_model=llm_model,
        system_prompt="你是一个严谨、克制、重视证据边界的生物医药方向研究员，并擅长从论文页面截图中读取图表和方法信息。",
        user_prompt=prompt,
        images=images,
    )


def simplify_markdown_response(markdown: str, max_chars: int = 9000) -> str:
    lines = []
    skip_prefixes = (
        "好的",
        "当然",
        "作为",
        "我将",
        "下面",
        "以下",
        "请注意",
    )
    for raw_line in str(markdown or "").splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            if lines and lines[-1] != "":
                lines.append("")
            continue
        stripped = line.strip()
        if not lines and any(stripped.startswith(prefix) for prefix in skip_prefixes):
            continue
        lines.append(line)
    cleaned = "\n".join(lines).strip()
    return compact_text(cleaned, max_chars)


def repair_markdown_tables(markdown: str) -> str:
    lines = str(markdown or "").splitlines()
    repaired: List[str] = []
    table_header_pattern = re.compile(r"^\s*\|\s*年份\s*\|.+\|\s*$")
    separator_pattern = re.compile(r"^\s*\|?\s*:?-{3,}:?\s*(\|\s*:?-{3,}:?\s*)+\|?\s*$")

    for index, line in enumerate(lines):
        repaired.append(line)
        if table_header_pattern.match(line):
            next_line = lines[index + 1] if index + 1 < len(lines) else ""
            if not separator_pattern.match(next_line):
                column_count = max(2, line.count("|") - 1)
                repaired.append("|" + "|".join(["---"] * column_count) + "|")
    return "\n".join(repaired)


def polish_breakthrough_report(markdown: str) -> str:
    cleaned = simplify_markdown_response(markdown, 12000)
    cleaned = repair_markdown_tables(cleaned)
    if "｜" in cleaned:
        cleaned = cleaned.replace("｜", "|")
    return cleaned.strip()


def demote_large_markdown_headings(markdown: str) -> str:
    lines = []
    for line in str(markdown or "").splitlines():
        stripped = line.lstrip()
        if stripped.startswith("# "):
            lines.append("#### " + stripped[2:].strip())
        elif stripped.startswith("## "):
            lines.append("#### " + stripped[3:].strip())
        else:
            lines.append(line)
    return "\n".join(lines)


def extract_markdown_table_rows(markdown: str) -> List[Dict[str, str]]:
    text = str(markdown or "").replace("｜", "|")
    normalized_lines = []
    for raw_line in text.splitlines():
        line = raw_line.strip()
        if "| 年份 |" in line and not line.startswith("|"):
            line = line[line.find("| 年份 |") :].strip()
        if line.startswith("|") and not line.endswith("|"):
            line = f"{line}|"
        if line.startswith("|") and line.endswith("|"):
            normalized_lines.append(line)
    lines = normalized_lines
    if not lines:
        return []

    header_index = -1
    for idx, line in enumerate(lines):
        if "年份" in line and ("里程碑" in line or "技术突破" in line or "代表文献" in line):
            header_index = idx
            break
    if header_index < 0:
        return []

    headers = [cell.strip() for cell in lines[header_index].strip("|").split("|")]
    rows: List[Dict[str, str]] = []
    for line in lines[header_index + 1 :]:
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        if not cells or all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        if len(cells) < len(headers):
            cells.extend([""] * (len(headers) - len(cells)))
        row = {headers[i]: cells[i] for i in range(min(len(headers), len(cells)))}
        if any(value for value in row.values()):
            rows.append(row)
    return rows


def normalize_breakthrough_rows(rows: Any) -> List[Dict[str, str]]:
    normalized: List[Dict[str, str]] = []
    if not isinstance(rows, list):
        return normalized

    key_map = {
        "年份": ["年份", "year"],
        "里程碑/技术突破": ["里程碑/技术突破", "里程碑", "技术突破", "milestone", "breakthrough"],
        "代表文献": ["代表文献", "representative_paper", "paper", "reference"],
        "DOI": ["DOI", "doi"],
        "关键意义": ["关键意义", "significance", "meaning", "impact"],
    }

    for item in rows:
        if not isinstance(item, dict):
            continue
        row: Dict[str, str] = {}
        for target, candidates in key_map.items():
            value = ""
            for key in candidates:
                if key in item and item.get(key) is not None:
                    value = strip_markup(item.get(key))
                    break
            row[target] = value or ("未检出" if target == "DOI" else "")
        if any(row.get(col) for col in BREAKTHROUGH_COLUMNS):
            normalized.append(row)

    def year_key(row: Dict[str, str]) -> tuple[int, str]:
        match = re.search(r"\d{4}", row.get("年份", ""))
        return (int(match.group(0)) if match else 9999, row.get("年份", ""))

    return sorted(normalized, key=year_key)


def normalize_breakthrough_report(data: Dict[str, Any]) -> Dict[str, Any]:
    rows = (
        data.get("milestones")
        or data.get("breakthroughs")
        or data.get("关键里程碑文献")
        or data.get("rows")
        or []
    )
    keywords = data.get("core_keywords") or data.get("核心关键词") or []
    if isinstance(keywords, str):
        keywords = [item.strip() for item in re.split(r"[;；,，、]", keywords) if item.strip()]

    return {
        "field_direction": strip_markup(data.get("field_direction") or data.get("领域方向")),
        "direction_overview": strip_markup(data.get("direction_overview") or data.get("方向概述")),
        "core_keywords": [strip_markup(item) for item in keywords if strip_markup(item)],
        "milestones": normalize_breakthrough_rows(rows),
        "usage_note": strip_markup(data.get("usage_note") or data.get("使用提醒") or "DOI 需用 CrossRef/PubMed/出版社页面二次核验。"),
    }


def breakthrough_report_to_dataframe(report: Dict[str, Any]) -> pd.DataFrame:
    rows = report.get("milestones") or []
    return pd.DataFrame(rows, columns=BREAKTHROUGH_COLUMNS)


def set_docx_run_font(run: Any, size: int = 11, bold: bool = False, color: Optional[str] = None) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Pt, RGBColor

    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_docx_paragraph(doc: Any, text: str, *, size: int = 11, bold: bool = False, color: Optional[str] = None) -> Any:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = 6
    paragraph.paragraph_format.line_spacing = 1.15
    run = paragraph.add_run(strip_markup(text))
    set_docx_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def add_docx_heading(doc: Any, text: str, level: int = 1) -> Any:
    size = 16 if level == 1 else 13
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = 12 if level == 1 else 8
    paragraph.paragraph_format.space_after = 6
    run = paragraph.add_run(strip_markup(text))
    set_docx_run_font(run, size=size, bold=True, color="2E74B5" if level == 1 else "1F4D78")
    return paragraph


def add_markdownish_text_to_docx(doc: Any, text: str) -> None:
    for raw in str(text or "").splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("【") and "】" in line:
            label, body = line.split("】", 1)
            add_docx_heading(doc, label.strip("【】"), level=2)
            if body.strip("：: "):
                add_docx_paragraph(doc, body.strip("：: "))
        elif line.startswith(("-", "•")):
            paragraph = doc.add_paragraph(style=None)
            paragraph.paragraph_format.left_indent = None
            paragraph.paragraph_format.space_after = 4
            run = paragraph.add_run(line.lstrip("-• ").strip())
            set_docx_run_font(run)
        else:
            add_docx_paragraph(doc, line)


def set_table_cell_text(cell: Any, text: str, *, bold: bool = False, fill: Optional[str] = None) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    if fill:
        tc_pr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), fill)
        tc_pr.append(shd)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = 0
    run = paragraph.add_run(strip_markup(text))
    set_docx_run_font(run, size=9, bold=bold)


def add_breakthrough_table_to_docx(doc: Any, report: Dict[str, Any]) -> None:
    from docx.shared import Inches

    df = breakthrough_report_to_dataframe(report)
    table = doc.add_table(rows=1, cols=len(BREAKTHROUGH_COLUMNS))
    table.style = "Table Grid"
    table.autofit = False
    widths = [0.55, 1.55, 1.65, 1.35, 1.4]

    for col_idx, column in enumerate(BREAKTHROUGH_COLUMNS):
        cell = table.rows[0].cells[col_idx]
        cell.width = Inches(widths[col_idx])
        set_table_cell_text(cell, column, bold=True, fill="E8EEF5")

    for _, row_data in df.iterrows():
        cells = table.add_row().cells
        for col_idx, column in enumerate(BREAKTHROUGH_COLUMNS):
            cells[col_idx].width = Inches(widths[col_idx])
            set_table_cell_text(cells[col_idx], str(row_data.get(column, "")))


def build_pdf_deep_reading_docx_bytes(
    step1_report: str,
    breakthrough_report: Dict[str, Any],
    *,
    source_file: str,
    analysis_source: str,
) -> bytes:
    from docx import Document
    from docx.enum.section import WD_SECTION
    from docx.shared import Inches

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = 3
    title_run = title.add_run("PDF 文献精读报告")
    set_docx_run_font(title_run, size=22, bold=True)
    add_docx_paragraph(doc, f"来源文件：{source_file}", size=10, color="555555")
    add_docx_paragraph(doc, f"读取方式：{analysis_source}", size=10, color="555555")

    add_docx_heading(doc, "单篇文献深度剖析", level=1)
    add_markdownish_text_to_docx(doc, step1_report)

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_docx_heading(doc, "领域突破总结", level=1)
    if breakthrough_report.get("field_direction"):
        add_docx_paragraph(doc, f"领域方向：{breakthrough_report.get('field_direction')}", bold=True)
    if breakthrough_report.get("direction_overview"):
        add_docx_paragraph(doc, f"方向概述：{breakthrough_report.get('direction_overview')}")
    keywords = breakthrough_report.get("core_keywords") or []
    if keywords:
        add_docx_paragraph(doc, "核心关键词：" + "；".join(keywords))

    add_docx_heading(doc, "关键里程碑文献", level=2)
    add_breakthrough_table_to_docx(doc, breakthrough_report)
    if breakthrough_report.get("usage_note"):
        add_docx_paragraph(doc, breakthrough_report.get("usage_note"), size=9, color="555555")

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


def render_breakthrough_result(report: Dict[str, Any]) -> None:
    st.markdown("### 领域突破表")
    if report.get("field_direction"):
        st.write(f"**领域方向：** {report.get('field_direction')}")
    if report.get("direction_overview"):
        st.write(f"**方向概述：** {report.get('direction_overview')}")
    keywords = report.get("core_keywords") or []
    if keywords:
        st.write("**核心关键词：** " + "；".join(keywords))

    df = breakthrough_report_to_dataframe(report)
    if df.empty:
        st.info("模型未能生成足够的里程碑条目，可重试或切换模型。")
    st.dataframe(df, use_container_width=True, hide_index=True)
    if report.get("usage_note"):
        st.caption(report.get("usage_note"))


def analyze_field_breakthroughs(
    step1_report: str,
    pdf_text: str,
    llm_api_key: str,
    llm_provider: str,
    llm_base_url: str,
    llm_model: str,
) -> Dict[str, Any]:
    prompt = f"""
请基于下面的文献剖析报告和论文片段，生成一份“近 10 年领域突破追踪”。

输出必须是严格 JSON，不要输出 Markdown、表格文本或解释。

JSON 结构：
{{
  "field_direction": "用 8-18 个字概括细分方向",
  "direction_overview": "不超过 80 字说明过去 10 年演进主线",
  "core_keywords": ["关键词1", "关键词2", "关键词3", "关键词4", "关键词5"],
  "milestones": [
    {{
      "年份": "2016",
      "里程碑/技术突破": "简短突破名",
      "代表文献": "第一作者 et al., 年份, 期刊",
      "DOI": "未检出",
      "关键意义": "不超过 80 字的关键意义"
    }}
  ],
  "usage_note": "DOI 需用 CrossRef/PubMed/出版社页面二次核验。"
}}

要求：
1. milestones 按年份升序，优先列 6-12 条最关键文献。
2. 每格内容尽量控制在 80 字以内。
3. DOI 无法确定时写“未检出”，严禁编造 DOI。
4. 不要把整个表格塞进一个字符串，milestones 必须是对象数组。

文献剖析报告：
{step1_report}

论文片段：
{compact_text(pdf_text, 6000)}
""".strip()
    result = call_llm(
        llm_api_key=llm_api_key,
        llm_provider=llm_provider,
        llm_base_url=llm_base_url,
        llm_model=llm_model,
        system_prompt="你是一个熟悉生物医药技术史、酶工程、分子检测和基因组技术的专家。",
        user_prompt=prompt,
    )
    try:
        return normalize_breakthrough_report(parse_json_object(result))
    except Exception:
        rows = extract_markdown_table_rows(result)
        return normalize_breakthrough_report(
            {
                "field_direction": "",
                "direction_overview": simplify_markdown_response(result, 500),
                "core_keywords": [],
                "milestones": rows,
                "usage_note": "模型未返回 JSON，已尝试从文本中恢复表格；DOI 仍需二次核验。",
            }
        )


def render_pdf_deep_reading_tab(llm_api_key: str, llm_provider: str, llm_base_url: str, llm_model: str) -> None:
    st.subheader("PDF精读")
    st.caption("上传一篇论文，自动提炼研究亮点、实验路线与领域突破脉络。")

    with st.expander("使用说明", expanded=False):
        st.markdown(
            "- 快速模式读取核心页；增强模式优先使用 PyMuPDF4LLM，缺失时降级为 pypdf 全文抽取。\n"
            "- 结构化模式优先使用 Docling，缺失时降级为章节启发式抽取。\n"
            "- 图片读取适合扫描版或图表较多的 PDF，需要选择 Google Gemini。"
        )

    parse_mode = st.segmented_control(
        "解析模式",
        [PDF_PARSE_FAST, PDF_PARSE_ENHANCED, PDF_PARSE_STRUCTURED],
        default=PDF_PARSE_ENHANCED,
    )
    read_mode = st.radio(
        "PDF 读取方式",
        [PDF_TEXT_MODE, PDF_IMAGE_MODE, PDF_HYBRID_MODE],
        horizontal=True,
        help="图片模式适合扫描版、图表较多或文字提取乱码的 PDF；503 仍属于模型服务拥堵，需要重试或切换模型。",
    )
    if read_mode == PDF_IMAGE_MODE:
        st.caption("图片模式会把核心页面截图发送给 Gemini，能读取扫描版和图表，但更耗模型资源。")
    elif read_mode == PDF_HYBRID_MODE:
        st.caption("混合模式会先尝试文本提取；若文本为空、过短或解析失败，会自动转为 Gemini 图片识别。")

    uploaded_pdf = st.file_uploader("上传 PDF 文献", type=["pdf"])
    start_pdf = st.button("开始智能分析", type="primary", use_container_width=True)

    if start_pdf:
        if not validate_llm_config(llm_api_key, llm_base_url, llm_model):
            return
        if uploaded_pdf is None:
            st.warning("请先上传一个 PDF 文件。")
            return
        if read_mode in {PDF_IMAGE_MODE, PDF_HYBRID_MODE} and llm_provider != "Google Gemini":
            st.warning("当前图片识别模式请切换到 Google Gemini。OpenAI 兼容接口的视觉输入格式不统一，首版暂不启用。")
            return

        try:
            pdf_bytes = uploaded_pdf.getvalue()
            pdf_text = ""
            pdf_content: Dict[str, Any] = {}
            page_images: List[Dict[str, Any]] = []
            analysis_source = f"{parse_mode} / 文本读取"

            if read_mode == PDF_TEXT_MODE:
                with st.spinner("正在解析 PDF 内容..."):
                    pdf_content = extract_pdf_content(pdf_bytes, parse_mode)
                    pdf_text = str(pdf_content.get("markdown") or "")
                st.caption(f"解析器：{pdf_content.get('parser')}；证据等级：{pdf_content.get('evidence_level')}；状态：{pdf_content.get('status')}")
                with st.expander("已提取内容预览", expanded=False):
                    st.text_area("内容预览", pdf_content_to_prompt_text(pdf_content, 9000), height=260)
            elif read_mode == PDF_IMAGE_MODE:
                analysis_source = "图片模式"
                with st.spinner("正在将 PDF 核心页面渲染为图片..."):
                    page_images = render_core_pdf_pages_as_images(pdf_bytes)
                st.caption(f"已渲染 {len(page_images)} 张核心页面图片，准备交给 Gemini 识别。")
            else:
                with st.spinner("正在优先尝试解析 PDF 内容..."):
                    try:
                        pdf_content = extract_pdf_content(pdf_bytes, parse_mode)
                        pdf_text = str(pdf_content.get("markdown") or "")
                    except Exception as exc:
                        st.info(f"文本解析不可用，自动切换图片模式：{exc}")
                        pdf_text = ""
                if len(pdf_text) >= PDF_HYBRID_MIN_TEXT_CHARS:
                    analysis_source = f"混合模式：已使用 {parse_mode}"
                    st.caption(f"解析器：{pdf_content.get('parser')}；证据等级：{pdf_content.get('evidence_level')}；状态：{pdf_content.get('status')}")
                    with st.expander("已提取内容预览", expanded=False):
                        st.text_area("内容预览", pdf_content_to_prompt_text(pdf_content, 9000), height=260)
                else:
                    analysis_source = "混合模式：已自动转为图片识别"
                    if pdf_text:
                        st.info(f"文本内容较少（{len(pdf_text)} 字符），自动切换图片模式。")
                    with st.spinner("正在将 PDF 核心页面渲染为图片..."):
                        page_images = render_core_pdf_pages_as_images(pdf_bytes)
                    st.caption(f"已渲染 {len(page_images)} 张核心页面图片，准备交给 Gemini 识别。")

            with st.spinner("正在生成文献剖析..."):
                if page_images:
                    step1_report = analyze_pdf_deep_reading_with_images(
                        page_images,
                        llm_api_key.strip(),
                        llm_provider,
                        llm_base_url.strip(),
                        llm_model.strip(),
                    )
                elif pdf_content:
                    step1_report = analyze_pdf_content_deep_reading(
                        pdf_content,
                        llm_api_key.strip(),
                        llm_provider,
                        llm_base_url.strip(),
                        llm_model.strip(),
                    )
                else:
                    step1_report = analyze_pdf_deep_reading(
                        pdf_text,
                        llm_api_key.strip(),
                        llm_provider,
                        llm_base_url.strip(),
                        llm_model.strip(),
                    )

            with st.spinner("正在深挖近10年领域背景..."):
                step2_report = analyze_field_breakthroughs(
                    step1_report,
                    pdf_text,
                    llm_api_key.strip(),
                    llm_provider,
                    llm_base_url.strip(),
                    llm_model.strip(),
                )
        except Exception as exc:
            st.error(f"智能分析失败：{exc}")
            return

        st.success(f"智能分析完成。实际读取方式：{analysis_source}")
        try:
            word_bytes = build_pdf_deep_reading_docx_bytes(
                step1_report,
                step2_report,
                source_file=uploaded_pdf.name,
                analysis_source=analysis_source,
            )
            st.download_button(
                "下载 Word 精读报告",
                data=word_bytes,
                file_name="pdf_deep_reading_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
            )
        except Exception as exc:
            st.warning(f"Word 报告生成失败：{exc}")

        analysis_tab, breakthrough_tab = st.tabs(["单篇剖析", "领域突破"])
        with analysis_tab:
            st.markdown("### 单篇文献深度剖析")
            st.markdown(demote_large_markdown_headings(step1_report))
        with breakthrough_tab:
            render_breakthrough_result(step2_report)


def parse_json_object(text: str) -> Dict[str, Any]:
    raw = str(text or "").strip()
    raw = re.sub(r"^```(?:json)?\s*", "", raw, flags=re.I)
    raw = re.sub(r"\s*```$", "", raw)
    match = re.search(r"\{.*\}", raw, flags=re.S)
    if match:
        raw = match.group(0)
    try:
        data = json.loads(raw)
        return data if isinstance(data, dict) else {}
    except json.JSONDecodeError as exc:
        raise ApiError(f"模型没有返回可解析 JSON，请重试或切换模型。原始错误：{exc}") from exc


def normalize_ppt_analysis(data: Dict[str, Any], fallback_title: str) -> Dict[str, Any]:
    title = strip_markup(data.get("document_title")) or fallback_title
    journal_name = strip_markup(data.get("journal_name"))
    doi = normalize_doi(strip_markup(data.get("doi")))
    impact_factor = strip_markup(data.get("impact_factor"))
    impact_factor_note = strip_markup(data.get("impact_factor_note"))
    main_content = strip_markup(data.get("main_content")) or "模型未能提取核心内容。"
    framework = data.get("writing_framework") or []
    if not isinstance(framework, list):
        framework = [strip_markup(framework)]

    sections = []
    for item in data.get("main_content_sections") or []:
        if not isinstance(item, dict):
            continue
        sections.append(
            {
                "section_name": strip_markup(item.get("section_name")),
                "subtopic": strip_markup(item.get("subtopic")),
                "key_points": strip_markup(item.get("key_points")),
            }
        )

    figures = []
    for item in data.get("figures_analysis") or []:
        if not isinstance(item, dict):
            continue
        figures.append(
            {
                "figure_id": strip_markup(item.get("figure_id")) or f"fig{len(figures) + 1}",
                "content_summary": strip_markup(item.get("content_summary")),
                "design_purpose": strip_markup(item.get("design_purpose")),
            }
        )

    return {
        "document_title": title,
        "journal_name": journal_name,
        "doi": doi,
        "impact_factor": impact_factor,
        "impact_factor_note": impact_factor_note,
        "main_content": main_content,
        "writing_framework": [strip_markup(x) for x in framework if strip_markup(x)],
        "main_content_sections": sections,
        "figures_analysis": figures,
        "figure_table_legends": data.get("figure_table_legends") or [],
    }


def pdf_looks_image_only(pdf_bytes: bytes) -> bool:
    try:
        import fitz

        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        selected = selected_pdf_page_indexes(doc.page_count)
        text_chars = 0
        image_pages = 0
        checked = 0
        for page_index in selected:
            page = doc[page_index]
            text_chars += len((page.get_text() or "").strip())
            if page.get_images(full=True):
                image_pages += 1
            checked += 1
        return checked > 0 and text_chars < 30 and image_pages >= max(1, checked // 2)
    except Exception:
        return False


def patent_or_pdf_image_prompt(file_name: str, pages: List[int]) -> str:
    page_list = "、".join(str(page) for page in pages)
    return f"""
请阅读下面的 PDF 页面截图，文件名为 {file_name}，页面范围为第 {page_list} 页。
该文件可能是扫描版专利或图片型论文，PDF 没有可提取文本层。请根据页面截图进行 OCR 式识别和技术归纳，输出严格 JSON，不要输出 Markdown 或解释文字。

JSON 字段：
{{
  "document_title": "文献、专利或申请的完整标题；无法识别时写文件名",
  "journal_name": "期刊/来源/专利公开机构；专利可写 WIPO/PCT/Google Patents 等",
  "doi": "文献 DOI；专利或无法识别写空字符串，不要编造",
  "main_content": "2-4 句核心内容总结，说明技术问题、技术方案和用途",
  "writing_framework": ["背景与痛点", "核心方案", "关键权利要求/实施例", "验证路线", "应用"],
  "main_content_sections": [
    {{"section_name": "核心内容", "subtopic": "", "key_points": "适合放入 PPT 表格的一段话"}},
    {{"section_name": "背景与痛点", "subtopic": "", "key_points": "..."}},
    {{"section_name": "核心方案", "subtopic": "", "key_points": "..."}},
    {{"section_name": "关键权利要求/实施例", "subtopic": "", "key_points": "..."}},
    {{"section_name": "应用", "subtopic": "", "key_points": "..."}},
    {{"section_name": "局限与待核验信息", "subtopic": "", "key_points": "..."}}
  ],
  "figures_analysis": [
    {{"figure_id": "Page 1", "content_summary": "该页展示的代表内容", "design_purpose": "该页/图示用于说明什么"}}
  ]
}}

要求：
1. 不确定的申请号、标题、申请人、权利要求或数值必须写“未在截图中明确给出”，不要编造。
2. 如果截图里是专利首页，请优先提取标题、公开信息、摘要和技术领域。
3. figures_analysis 最多 6 项，使用 Page/Fig 编号均可。
""".strip()


def analyze_pdf_images_for_ppt(
    pdf_bytes: bytes,
    file_name: str,
    llm_api_key: str,
    llm_provider: str,
    llm_base_url: str,
    llm_model: str,
) -> Dict[str, Any]:
    if llm_provider != "Google Gemini":
        raise ApiError(
            "该 PDF 是图片型/扫描版，当前文本解析器无法读取文字。"
            "请在左侧把 LLM 服务商切换为 Google Gemini 后重试，或先对 PDF 做 OCR 再上传。"
        )
    images = render_core_pdf_pages_as_images(pdf_bytes)
    if not images:
        raise ApiError("扫描版 PDF 未能渲染为页面图片，无法进行图片识别。")

    prompt = patent_or_pdf_image_prompt(file_name, [image["page"] for image in images])
    result = call_gemini_llm_with_images(
        llm_api_key=llm_api_key,
        llm_base_url=llm_base_url,
        llm_model=llm_model,
        system_prompt="你是擅长专利、论文页面 OCR 识别和技术汇报结构化的学术报告设计师。",
        user_prompt=prompt,
        images=images,
        timeout=240,
    )
    analysis = normalize_ppt_analysis(parse_json_object(result), fallback_title=file_name)
    analysis["doi"] = analysis.get("doi") or "待核验"
    analysis["impact_factor"] = "不适用" if "patent" in file_name.lower() or file_name.upper().startswith("WO") else "待核验"
    analysis["impact_factor_note"] = "专利/扫描件"
    analysis["figure_table_legends"] = []
    return analysis


def analyze_pdf_for_ppt(
    pdf_bytes: bytes,
    file_name: str,
    llm_api_key: str,
    llm_provider: str,
    llm_base_url: str,
    llm_model: str,
    parse_mode: str,
) -> Dict[str, Any]:
    if pdf_looks_image_only(pdf_bytes):
        return analyze_pdf_images_for_ppt(
            pdf_bytes=pdf_bytes,
            file_name=file_name,
            llm_api_key=llm_api_key,
            llm_provider=llm_provider,
            llm_base_url=llm_base_url,
            llm_model=llm_model,
        )

    try:
        content = extract_pdf_content(pdf_bytes, parse_mode)
    except ApiError as exc:
        message = str(exc)
        if "未能从 PDF 中提取到文本" in message or "未能从 PDF 中提取到文本" in message or "未能从 PDF 中提取到" in message:
            return analyze_pdf_images_for_ppt(
                pdf_bytes=pdf_bytes,
                file_name=file_name,
                llm_api_key=llm_api_key,
                llm_provider=llm_provider,
                llm_base_url=llm_base_url,
                llm_model=llm_model,
            )
        raise
    prompt_text = pdf_content_to_prompt_text(content, PPT_ANALYSIS_MAX_CHARS)
    extracted_doi = content.get("doi") or ""
    figure_table_legends = content.get("figure_table_legends") or []
    system_prompt = "你是擅长科研文献汇报、专利技术拆解和中文 PPT 结构化表达的学术报告设计师。"
    user_prompt = f"""
请阅读以下 PDF 文献/专利内容，输出严格 JSON，不要输出 Markdown 或解释文字。

JSON 字段：
{{
  "document_title": "文献或专利完整标题",
  "journal_name": "期刊或来源名称；专利则写专利来源/公开机构，无法识别写空字符串",
  "doi": "文献 DOI；没有 DOI 或无法识别写空字符串，不要编造",
  "main_content": "2-4 句核心内容总结，说明技术亮点、痛点和结论",
  "writing_framework": ["背景与痛点", "核心原理", "验证路线", "应用", "总结讨论"],
  "main_content_sections": [
    {{"section_name": "核心内容", "subtopic": "", "key_points": "适合放入 PPT 表格的一段话"}},
    {{"section_name": "背景与痛点", "subtopic": "", "key_points": "..."}},
    {{"section_name": "核心原理与概念验证", "subtopic": "", "key_points": "..."}},
    {{"section_name": "平台/复杂度验证", "subtopic": "", "key_points": "..."}},
    {{"section_name": "应用", "subtopic": "可选二级主题", "key_points": "..."}},
    {{"section_name": "总结与讨论", "subtopic": "", "key_points": "..."}},
    {{"section_name": "实验与分析方法", "subtopic": "", "key_points": "..."}}
  ],
  "figures_analysis": [
    {{"figure_id": "Fig. 1", "content_summary": "依据图注概括图中展示的代表内容", "design_purpose": "结合正文和图注解释图的设计目的"}},
    {{"figure_id": "Fig. 2", "content_summary": "依据图注概括图中展示的代表内容", "design_purpose": "结合正文和图注解释图的设计目的"}}
  ]
}}

要求：
1. main_content_sections 尽量生成 5-7 行，文字短、信息密度高，适合蓝色表格 PPT。
2. figures_analysis 必须优先使用“全文图例/图注”中的图表编号和图注内容，不要把整页 PDF 截图当作图表内容。
3. figures_analysis 只保留最关键的图/表，最多 8 项；如果图注里有 Fig. 1、Fig. 2、Table 1 等，请保持这些编号。
4. 不确定的内容要写“未在文本中明确给出”，不要编造 DOI、数值或实验结论。
5. document_title 必须是文章完整英文题名，不要只写短标题；doi 只能使用 PDF 文本中明确出现的 DOI。

文件名：{file_name}

PDF 内容：
{prompt_text}
""".strip()
    result = call_llm(
        llm_api_key=llm_api_key,
        llm_provider=llm_provider,
        llm_base_url=llm_base_url,
        llm_model=llm_model,
        system_prompt=system_prompt,
        user_prompt=user_prompt,
        timeout=180,
    )
    analysis = normalize_ppt_analysis(parse_json_object(result), fallback_title=file_name)
    if extracted_doi:
        analysis["doi"] = extracted_doi
    if not analysis.get("doi"):
        analysis["doi"] = "待核验"
    impact = lookup_impact_factor(analysis.get("journal_name"))
    analysis["impact_factor"] = impact["if"]
    analysis["impact_factor_note"] = impact["note"]
    analysis["figure_table_legends"] = figure_table_legends
    return analysis


def render_ppt_report_tab(llm_api_key: str, llm_provider: str, llm_base_url: str, llm_model: str) -> None:
    st.subheader("PPT汇报")
    st.caption("上传多篇论文/专利 PDF，自动生成同一个 PowerPoint，并在首页汇总全部附件。")

    with st.container(border=True):
        col_a, col_b, col_c = st.columns([1.2, 0.9, 0.9])
        with col_a:
            uploaded_files = st.file_uploader(
                "上传 PDF（可多选）",
                type=["pdf"],
                accept_multiple_files=True,
                key="ppt_pdf_uploads",
            )
        with col_b:
            parse_mode = st.selectbox(
                "PDF 解析模式",
                [PDF_PARSE_ENHANCED, PDF_PARSE_FAST, PDF_PARSE_STRUCTURED],
                key="ppt_parse_mode",
            )
        with col_c:
            output_name = st.text_input(
                "输出文件名",
                value="Academic_Analysis_Report.pptx",
                key="ppt_output_name",
            )
            if not output_name.lower().endswith(".pptx"):
                output_name += ".pptx"

    st.info("版式为内置 PPT 格式：第一页为上传附件汇总表，后续页面按文献展示写作主要内容和关键图表/图注解析。")
    start = st.button("生成 PPT", type="primary", use_container_width=True, key="ppt_generate")

    if not start:
        return
    if not validate_llm_config(llm_api_key, llm_base_url, llm_model):
        return
    if not uploaded_files:
        st.warning("请先上传至少一个 PDF。")
        return

    analyses: List[Dict[str, Any]] = []
    progress = st.progress(0.0, text="准备开始...")
    log_box = st.empty()

    for idx, uploaded in enumerate(uploaded_files, start=1):
        pdf_bytes = uploaded.getvalue()
        log_box.info(f"正在分析 {idx}/{len(uploaded_files)}：{uploaded.name}")
        try:
            analysis = analyze_pdf_for_ppt(
                pdf_bytes=pdf_bytes,
                file_name=uploaded.name,
                llm_api_key=llm_api_key.strip(),
                llm_provider=llm_provider,
                llm_base_url=llm_base_url.strip(),
                llm_model=llm_model.strip(),
                parse_mode=parse_mode,
            )
            analysis["source_file"] = uploaded.name
            analyses.append(analysis)
            legend_count = len(analysis.get("figure_table_legends") or [])
            if legend_count:
                st.caption(f"{uploaded.name} 已读取到 {legend_count} 条图例/图注。")
            else:
                st.caption(f"{uploaded.name} 未从文本层读取到明确图例/图注，可尝试结构化模式或 OCR 后再上传。")
            progress.progress(idx / len(uploaded_files), text=f"已完成 {idx}/{len(uploaded_files)}")
        except Exception as exc:
            st.error(f"{uploaded.name} 解析失败：{exc}")

    if not analyses:
        st.error("没有成功解析任何 PDF，无法生成 PPT。")
        return

    with st.spinner("正在排版并生成 PPT..."):
        pptx_bytes = build_pptx_bytes(analyses)
    log_box.success("PPT 已生成。")
    st.download_button(
        "下载 PPT",
        data=pptx_bytes,
        file_name=output_name,
        mime="application/vnd.openxmlformats-officedocument.presentationml.presentation",
        use_container_width=True,
    )

    with st.expander("本次解析标题", expanded=False):
        for idx, analysis in enumerate(analyses, start=1):
            st.write(f"{idx}. {analysis.get('document_title')}")


LITERATURE_SOURCE_LABELS = {
    "pubmed": "PubMed",
    "europepmc": "Europe PMC",
    "openalex": "OpenAlex",
    "crossref": "Crossref",
}


def literature_record_key(record: PaperRecord) -> str:
    for prefix, value in (
        ("doi", record.doi),
        ("pmid", record.pmid),
        ("pmcid", record.pmcid),
        ("openalex", record.openalex_id),
    ):
        if value:
            return f"{prefix}:{value.strip().lower()}"
    fallback = "\0".join((record.title, str(record.year or ""), record.authors[0] if record.authors else ""))
    return f"title:{hashlib.sha256(fallback.encode('utf-8')).hexdigest()}"


def literature_quality_label(record: PaperRecord) -> str:
    if record.is_retracted:
        return "已撤稿"
    if record.summary_ready and record.reference_ready:
        return "摘要/引用完整"
    if record.summary_ready:
        return "可总结"
    if record.reference_ready:
        return "可引用"
    return "需补全"


def render_literature_search_tab(
    llm_api_key: str,
    llm_provider: str,
    llm_base_url: str,
    llm_model: str,
    openalex_api_key: str,
    email: str,
) -> None:
    st.subheader("文献检索")
    st.caption("并行检索 PubMed、Europe PMC、OpenAlex 和 Crossref，统一去重后导出引用记录。")

    has_openalex_key = bool(openalex_api_key.strip())
    current_year = datetime.now().year
    with st.form("literature_search_form"):
        query = st.text_area(
            "研究问题或关键词",
            placeholder="例如：CRISPR base editing 在遗传病治疗中的进展",
            height=100,
        )
        search_attachments = st.file_uploader(
            "上传研究材料（可选）",
            type=SUPPORTED_ATTACHMENT_TYPES,
            accept_multiple_files=True,
            key="literature_search_attachments",
            help="上传论文、研究计划或笔记后，LLM 会结合材料生成分库检索式。",
        )
        st.markdown("**检索来源**")
        source_columns = st.columns(4)
        use_pubmed = source_columns[0].checkbox("PubMed", value=True)
        use_europepmc = source_columns[1].checkbox("Europe PMC", value=True)
        use_openalex = source_columns[2].checkbox(
            "OpenAlex",
            value=has_openalex_key,
            disabled=not has_openalex_key,
        )
        use_crossref = source_columns[3].checkbox("Crossref", value=True)
        if not has_openalex_key:
            st.caption("OpenAlex 需要 API Key；请在左侧高级设置或部署 Secrets 中配置后启用。")

        setting_columns = st.columns(3)
        start_year = int(setting_columns[0].number_input(
            "起始年份", min_value=1800, max_value=current_year + 1,
            value=max(1800, current_year - 5), step=1,
        ))
        end_year = int(setting_columns[1].number_input(
            "结束年份", min_value=1800, max_value=current_year + 1,
            value=current_year, step=1,
        ))
        per_source_limit = int(setting_columns[2].number_input(
            "每个来源返回数", min_value=1, max_value=50, value=20, step=1,
        ))
        smart_rewrite = st.checkbox(
            "使用 LLM 智能改写分库检索式",
            value=bool(llm_api_key.strip()),
            disabled=not bool(llm_api_key.strip()),
            help="仅调用一次 LLM；失败时自动使用原始关键词。",
        )
        submitted = st.form_submit_button("开始检索", type="primary", use_container_width=True)

    if submitted:
        attachment_context, attachment_warnings = build_attachment_context(search_attachments, max_chars=8_000)
        for warning in attachment_warnings:
            st.warning(warning)
        effective_query = query.strip()
        if attachment_context and smart_rewrite:
            effective_query = (
                f"{effective_query or '请根据附件提炼研究问题并设计检索式。'}\n\n"
                f"以下为用户上传的研究材料：\n{attachment_context}"
            )
        selected_sources = tuple(
            source for source, enabled in (
                ("pubmed", use_pubmed),
                ("europepmc", use_europepmc),
                ("openalex", use_openalex and has_openalex_key),
                ("crossref", use_crossref),
            ) if enabled
        )
        if not effective_query:
            st.warning("请输入研究问题或上传可读取的研究材料。")
        elif attachment_context and not smart_rewrite and not query.strip():
            st.warning("仅使用附件发起检索需要 LLM API Key；也可以手动填写关键词后直接检索。")
        elif not selected_sources:
            st.warning("请至少选择一个检索来源。")
        elif start_year > end_year:
            st.warning("起始年份不能晚于结束年份。")
        elif any(source in selected_sources for source in ("pubmed", "crossref")) and not email.strip():
            st.warning("请在左侧高级设置填写联系邮箱，以符合 NCBI/Crossref 的 API 使用规范。")
        else:
            credentials = {
                "openalex_api_key": openalex_api_key.strip(),
                "ncbi_api_key": config_value("NCBI_API_KEY").strip(),
                "email": email.strip(),
                "ncbi_tool": "biomedical-literature-workbench",
            }
            credential_marker = hashlib.sha256(
                "\0".join((llm_api_key, openalex_api_key, credentials["ncbi_api_key"])).encode("utf-8")
            ).hexdigest()
            if st.session_state.get("literature_credential_marker") != credential_marker:
                st.session_state["literature_cache_scope"] = os.urandom(16).hex()
                st.session_state["literature_credential_marker"] = credential_marker
            cache_scope = st.session_state.setdefault("literature_cache_scope", os.urandom(16).hex())
            try:
                with st.spinner("正在并行检索并合并结果..."):
                    result = cached_literature_search(
                        query=effective_query,
                        selected_sources=selected_sources,
                        start_year=start_year,
                        end_year=end_year,
                        limit=per_source_limit,
                        smart_rewrite=smart_rewrite,
                        llm_provider=llm_provider,
                        llm_base_url=llm_base_url,
                        llm_model=llm_model,
                        cache_scope=cache_scope,
                        _credentials=credentials,
                        _llm_api_key=llm_api_key,
                    )
                st.session_state["literature_search_result"] = result
                st.session_state["literature_selected_keys"] = set()
                for widget_key in (
                    "literature_filter_sources", "literature_filter_years",
                    "literature_filter_oa", "literature_filter_summary",
                    "literature_sort",
                ):
                    st.session_state.pop(widget_key, None)
            except Exception as exc:
                st.error(f"检索失败：{exc}")

    result = st.session_state.get("literature_search_result")
    if not isinstance(result, FederatedSearchResult):
        st.info("填写研究问题并开始检索。无 LLM Key 时会直接使用原始关键词。")
        return

    records = list(result.records)
    summary_count = sum(record.summary_ready for record in records)
    oa_count = sum(bool(record.is_open_access and record.oa_url) for record in records)
    metric_columns = st.columns(6)
    metric_columns[0].metric("原始记录", result.raw_count)
    metric_columns[1].metric("唯一文献", result.unique_count)
    metric_columns[2].metric("去除重复", result.duplicate_count)
    metric_columns[3].metric("可总结", summary_count)
    metric_columns[4].metric("开放获取", oa_count)
    metric_columns[5].metric("耗时", f"{result.elapsed_seconds:.1f}s")

    source_metric_columns = st.columns(4)
    for column, source in zip(source_metric_columns, LITERATURE_SOURCE_LABELS):
        column.metric(LITERATURE_SOURCE_LABELS[source], result.source_counts.get(source, 0))

    with st.expander("分库检索式与来源状态", expanded=bool(result.errors)):
        for source, source_query in result.source_queries.items():
            st.markdown(f"**{LITERATURE_SOURCE_LABELS.get(source, source)}**")
            st.code(source_query, language=None)
        if result.terms_en or result.terms_zh:
            st.caption(
                f"英文扩展词：{'; '.join(result.terms_en) or '无'} | "
                f"中文扩展词：{'; '.join(result.terms_zh) or '无'}"
            )
        if result.errors:
            st.markdown("**来源错误（其他来源结果仍保留）**")
            for source, error in result.errors.items():
                st.warning(f"{LITERATURE_SOURCE_LABELS.get(source, source)}：{error}")
        else:
            st.success("已完成全部所选来源。")

    if not records:
        st.warning("当前检索没有返回文献，请调整关键词、年份或来源后重试。")
        return

    st.markdown("### 筛选与排序")
    source_options = [
        source for source in LITERATURE_SOURCE_LABELS
        if any(source in (record.sources or [record.source]) for record in records)
    ]
    filter_columns = st.columns(4)
    source_filter = filter_columns[0].multiselect(
        "来源", source_options, default=source_options,
        format_func=lambda value: LITERATURE_SOURCE_LABELS.get(value, value),
        key="literature_filter_sources",
    )
    known_years = [record.year for record in records if record.year is not None]
    if known_years and min(known_years) < max(known_years):
        year_bounds = (min(known_years), max(known_years))
        selected_year_range = filter_columns[1].slider(
            "年份", min_value=year_bounds[0], max_value=year_bounds[1],
            value=year_bounds, key="literature_filter_years",
        )
    else:
        only_year = known_years[0] if known_years else None
        filter_columns[1].text_input("年份", value=str(only_year or "未知"), disabled=True)
        year_bounds = (only_year, only_year)
        selected_year_range = year_bounds
    oa_filter = filter_columns[2].selectbox(
        "开放获取", ["全部", "仅开放获取", "排除开放获取"],
        key="literature_filter_oa",
    )
    summary_filter = filter_columns[3].selectbox(
        "摘要状态", ["全部", "可总结", "摘要不足"],
        key="literature_filter_summary",
    )
    sort_mode = st.segmented_control(
        "排序", ["综合相关性", "年份（新到旧）", "引用数（高到低）"],
        default="综合相关性", key="literature_sort",
    )

    selected_source_set = set(source_filter)
    filtered_records = [
        record for record in records
        if selected_source_set.intersection(record.sources or [record.source])
    ]
    if known_years and selected_year_range[0] is not None:
        full_year_range = selected_year_range == year_bounds
        filtered_records = [
            record for record in filtered_records
            if (record.year is None and full_year_range)
            or (record.year is not None and selected_year_range[0] <= record.year <= selected_year_range[1])
        ]
    if oa_filter == "仅开放获取":
        filtered_records = [record for record in filtered_records if record.is_open_access and record.oa_url]
    elif oa_filter == "排除开放获取":
        filtered_records = [record for record in filtered_records if not (record.is_open_access and record.oa_url)]
    if summary_filter == "可总结":
        filtered_records = [record for record in filtered_records if record.summary_ready]
    elif summary_filter == "摘要不足":
        filtered_records = [record for record in filtered_records if not record.summary_ready]
    if sort_mode == "年份（新到旧）":
        filtered_records.sort(key=lambda record: (record.year or 0, record.score), reverse=True)
    elif sort_mode == "引用数（高到低）":
        filtered_records.sort(key=lambda record: (record.citations or -1, record.score), reverse=True)
    else:
        filtered_records.sort(key=lambda record: record.rank or 999999)

    selected_keys = set(st.session_state.get("literature_selected_keys", set()))
    table_rows = []
    for record in filtered_records:
        record_key = literature_record_key(record)
        authors = "; ".join(record.authors[:3])
        if len(record.authors) > 3:
            authors += " 等"
        identifiers = record.doi or (f"PMID:{record.pmid}" if record.pmid else "")
        table_rows.append({
            "选择": record_key in selected_keys,
            "题名": f"[撤稿] {record.title}" if record.is_retracted else record.title,
            "作者": authors,
            "年份": record.year,
            "期刊": record.publication,
            "DOI / PMID": identifiers,
            "来源": "; ".join(LITERATURE_SOURCE_LABELS.get(source, source) for source in record.sources),
            "引用数": record.citations,
            "质量": literature_quality_label(record),
            "链接": record.oa_url or record.url,
            "_record_key": record_key,
        })

    st.markdown("### 文献列表")
    st.caption(f"当前筛选 {len(filtered_records)} 篇；勾选后优先导出勾选项，未勾选则导出当前筛选结果。")
    if table_rows:
        table_frame = pd.DataFrame(table_rows)
        editor_signature = hashlib.sha256(json.dumps({
            "sources": source_filter,
            "years": selected_year_range,
            "oa": oa_filter,
            "summary": summary_filter,
            "sort": sort_mode,
        }, ensure_ascii=True, sort_keys=True).encode("utf-8")).hexdigest()[:16]
        edited_frame = st.data_editor(
            table_frame,
            key=f"literature_result_editor_{editor_signature}",
            use_container_width=True,
            hide_index=True,
            height=min(720, 38 + 36 * len(table_frame)),
            column_order=[
                "选择", "题名", "作者", "年份", "期刊", "DOI / PMID",
                "来源", "引用数", "质量", "链接",
            ],
            column_config={
                "选择": st.column_config.CheckboxColumn("选择", width="small"),
                "题名": st.column_config.TextColumn("题名", width="large"),
                "作者": st.column_config.TextColumn("作者", width="medium"),
                "年份": st.column_config.NumberColumn("年份", format="%d", width="small"),
                "引用数": st.column_config.NumberColumn("引用数", format="%d", width="small"),
                "链接": st.column_config.LinkColumn("链接", display_text="打开", width="small"),
            },
            disabled=[
                "题名", "作者", "年份", "期刊", "DOI / PMID", "来源",
                "引用数", "质量", "链接", "_record_key",
            ],
        )
        visible_keys = set(table_frame["_record_key"].tolist())
        selected_keys.difference_update(visible_keys)
        selected_keys.update(
            str(row["_record_key"])
            for row in edited_frame.to_dict("records")
            if row.get("选择")
        )
        st.session_state["literature_selected_keys"] = selected_keys
    else:
        st.info("没有符合当前筛选条件的文献。")

    selected_records = [
        record for record in records if literature_record_key(record) in selected_keys
    ]
    export_records = selected_records or filtered_records
    export_label = f"勾选的 {len(selected_records)} 篇" if selected_records else f"筛选后的 {len(filtered_records)} 篇"
    st.markdown("### 导出")
    st.caption(f"将导出{export_label}。Excel 始终包含 Papers、Source_Summary、Search_Queries 三个工作表。")
    if export_records:
        export_result = replace(result, records=export_records, unique_count=len(export_records))
        date_suffix = datetime.now().strftime("%Y%m%d")
        download_columns = st.columns(3)
        download_columns[0].download_button(
            "下载 Excel",
            data=export_excel(export_result),
            file_name=f"literature_search_{date_suffix}.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )
        download_columns[1].download_button(
            "下载 RIS",
            data=export_ris(export_records),
            file_name=f"literature_search_{date_suffix}.ris",
            mime="application/x-research-info-systems",
            use_container_width=True,
        )
        download_columns[2].download_button(
            "下载 BibTeX",
            data=export_bibtex(export_records),
            file_name=f"literature_search_{date_suffix}.bib",
            mime="application/x-bibtex",
            use_container_width=True,
        )

    fetched_dates = [record.fetched_at[:10] for record in records if record.fetched_at]
    fetched_date = max(fetched_dates, default=datetime.now(timezone.utc).date().isoformat())
    st.markdown(
        "<small>数据来源："
        '<a href="https://pubmed.ncbi.nlm.nih.gov/">PubMed/NCBI</a> · '
        '<a href="https://europepmc.org/">Europe PMC</a> · '
        '<a href="https://openalex.org/">OpenAlex</a> · '
        '<a href="https://www.crossref.org/">Crossref</a>'
        f"；抓取日期：{fetched_date}。仅提供来源页和上游确认的开放获取链接。</small>",
        unsafe_allow_html=True,
    )


def render_citation_tracking_tab(openalex_api_key: str, email: str, max_papers: int) -> None:
    st.subheader("引用追踪")
    st.caption("输入 DOI / OpenAlex Paper ID，或上传含有标识符的材料，抓取引用文献并导出 Excel。")
    with st.expander("使用说明", expanded=False):
        st.markdown(
            "- 数据源为 OpenAlex，不需要 Semantic Scholar API Key。\n"
            "- 结果为客观元数据：标题、年份、DOI、摘要、引用数、期刊和作者。\n"
            "- 开放全文上下文模式只处理 OA PDF，不绕过付费墙；没有全文时会降级为元数据证据。"
        )
    citation_attachments = st.file_uploader(
        "上传文献或标识符清单（可选）",
        type=SUPPORTED_ATTACHMENT_TYPES,
        accept_multiple_files=True,
        key="citation_attachments",
        help="系统会从附件文本中识别第一个 DOI 或 OpenAlex Paper ID。",
    )
    attachment_context, attachment_warnings = build_attachment_context(citation_attachments, max_chars=12_000)
    attachment_doi = extract_doi_from_text(attachment_context)
    openalex_match = re.search(r"(?i)\bW\d{6,}\b", attachment_context)
    attachment_identifier = attachment_doi or (openalex_match.group(0) if openalex_match else "")
    if attachment_identifier:
        st.caption(f"已从附件识别：`{attachment_identifier}`（可在下方手动覆盖）")
    identifier = st.text_input(
        "目标文献 DOI 或 OpenAlex Paper ID",
        placeholder="例如：10.1038/s41586-020-2649-2 或 W3035965352",
    )
    citation_mode = st.segmented_control(
        "分析模式",
        [CITATION_METADATA_ONLY, CITATION_CONTEXT_MODE],
        default=CITATION_METADATA_ONLY,
    )
    if citation_mode == CITATION_CONTEXT_MODE:
        st.caption("实验功能：会尝试下载开放获取 PDF 并定位目标 DOI 附近上下文，速度较慢且覆盖率有限。")
    start = st.button("开始分析", type="primary", use_container_width=True)

    if start:
        for warning in attachment_warnings:
            st.warning(warning)
        resolved_identifier = identifier.strip() or attachment_identifier
        if not resolved_identifier:
            st.warning("请输入目标文献的 DOI / Paper ID，或上传包含标识符的附件。")
            return

        progress = st.progress(0, text="准备开始...")
        with st.expander("实时分析日志", expanded=True):
            title_log = st.empty()

        try:
            with st.spinner("系统正在分析，请稍候..."):
                df = run_analysis(
                    identifier=resolved_identifier,
                    openalex_api_key=openalex_api_key.strip(),
                    email=email.strip(),
                    max_papers=int(max_papers),
                    citation_mode=citation_mode,
                    progress=progress,
                    title_log=title_log,
                )
        except Exception as exc:
            st.error(f"分析失败：{exc}")
            return

        st.subheader("结果概览")
        doi_count = int(df["DOI"].astype(bool).sum()) if not df.empty and "DOI" in df else 0
        years = pd.to_numeric(df["年份"], errors="coerce").dropna() if not df.empty and "年份" in df else pd.Series(dtype=float)
        year_range = "无" if years.empty else f"{int(years.min())}-{int(years.max())}"
        metric_cols = st.columns(3)
        metric_cols[0].metric("文献数", len(df))
        metric_cols[1].metric("含 DOI", doi_count)
        metric_cols[2].metric("年份范围", year_range)

        st.subheader("文献列表")
        st.dataframe(df, use_container_width=True, hide_index=True)

        excel_bytes = dataframe_to_excel_bytes(df)
        st.download_button(
            label="下载 Excel 分析报告",
            data=excel_bytes,
            file_name="citation_analysis_report.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            use_container_width=True,
        )


def render_research_tools_tab(
    llm_api_key: str,
    llm_provider: str,
    llm_base_url: str,
    llm_model: str,
    store: ResearchStore,
    identity: AuthIdentity,
    conversation_id: str,
) -> None:
    render_research_workspace(
        NATURE_TOOL_CONFIGS,
        llm_api_key,
        llm_provider,
        llm_base_url,
        llm_model,
        call_llm,
        store,
        identity.user_id,
        conversation_id,
    )


def render_app() -> None:
    st.set_page_config(
        page_title="生物医学科研工作台",
        page_icon="🧬",
        layout="wide",
        initial_sidebar_state="auto",
    )
    identity = require_authenticated_user()
    try:
        store = cached_research_store()
        store.ensure_profile(identity)
    except ConfigurationError as exc:
        st.error(f"云端数据服务配置失败：{exc}")
        st.info("生产环境默认要求 PostgreSQL + 私有 COS；仅本地开发可显式设置 APP_PERSISTENCE_BACKEND=memory。")
        st.stop()
        return
    except Exception as exc:
        st.error(f"云端数据服务暂不可用：{exc}")
        st.stop()
        return
    apply_page_style()
    st.markdown(
        """
        <div class="app-hero">
            <div class="app-kicker">Biomedical Research Workspace</div>
            <h1>生物医学科研工作台</h1>
            <p>检索、阅读、分析、写作与汇报集中在一个清晰工作区，并支持附件和连续对话。</p>
        </div>
        """,
        unsafe_allow_html=True,
    )

    with st.sidebar:
        st.header("工作区")
        workspace = st.radio(
            "选择功能",
            ["PDF精读", "文献检索", "引用追踪", "科研工具", "PPT汇报"],
            label_visibility="collapsed",
        )
        st.caption(f"已登录：{identity.email}")
        if st.button("退出登录", use_container_width=True):
            st.logout()
        st.divider()
        st.subheader("模型配置")
        llm_provider = st.selectbox("LLM 服务商", ["OpenAI 兼容", "Google Gemini"])
        defaults = load_config_defaults(llm_provider)
        llm_api_key = st.text_input("LLM API Key", type="password", key="browser_session_llm_api_key")

        with st.expander("高级设置", expanded=False):
            llm_base_url = st.text_input("Base URL", value=defaults["llm_base_url"])
            llm_model = st.text_input("模型名称", value=defaults["llm_model"])
            openalex_api_key = st.text_input("OpenAlex API Key（可选）", value=defaults["openalex_api_key"], type="password")
            email = st.text_input("OpenAlex Email", value=defaults["openalex_email"])
            max_papers = st.number_input("最多文献数", min_value=1, max_value=1000, value=20, step=1)

        st.caption("LLM Key 仅驻留当前浏览器会话，不写入数据库、日志或对象存储。")

        if st.button("测试连接", use_container_width=True):
            if validate_llm_config(llm_api_key, llm_base_url, llm_model):
                try:
                    with st.spinner("正在测试..."):
                        reply = call_llm(
                            llm_api_key=llm_api_key.strip(),
                            llm_provider=llm_provider,
                            llm_base_url=llm_base_url.strip(),
                            llm_model=llm_model.strip(),
                            system_prompt="你只需要回答连接正常。",
                            user_prompt="请回复：连接正常",
                            timeout=30,
                        )
                    st.success(f"连接成功：{reply[:60]}")
                except Exception as exc:
                    st.error(f"连接失败：{exc}")

        current_conversation_id = ""
        if workspace == "科研工具":
            st.divider()
            current_conversation_id = render_research_conversation_sidebar(store, identity.user_id)

    if workspace == "PDF精读":
        render_pdf_deep_reading_tab(llm_api_key, llm_provider, llm_base_url, llm_model)
    elif workspace == "文献检索":
        render_literature_search_tab(
            llm_api_key,
            llm_provider,
            llm_base_url,
            llm_model,
            openalex_api_key,
            email,
        )
    elif workspace == "引用追踪":
        render_citation_tracking_tab(openalex_api_key, email, int(max_papers))
    elif workspace == "科研工具":
        render_research_tools_tab(
            llm_api_key,
            llm_provider,
            llm_base_url,
            llm_model,
            store,
            identity,
            current_conversation_id,
        )
    elif workspace == "PPT汇报":
        render_ppt_report_tab(llm_api_key, llm_provider, llm_base_url, llm_model)


if __name__ == "__main__":
    render_app()
